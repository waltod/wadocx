"""
Tracked-changes (revision) tools for WaDocx MCP.

High-level MCP tools for authoring and resolving Word tracked changes: enabling
track-changes mode, inserting/deleting/replacing text as revisions, and
accepting or rejecting all changes.
"""
import os
from typing import Optional

from docx import Document

from word_document_server.utils.file_utils import check_file_writeable, ensure_docx_extension
from word_document_server.core import revisions as R


def _prepare(filename: str):
    """Validate the document path and writeability; return (filename, error)."""
    filename = ensure_docx_extension(filename)
    if not os.path.exists(filename):
        return filename, f"Document {filename} does not exist"
    is_writeable, error_message = check_file_writeable(filename)
    if not is_writeable:
        return filename, (
            f"Cannot modify document: {error_message}. Consider creating a copy first."
        )
    return filename, None


async def enable_track_changes(filename: str, enabled: bool = True) -> str:
    """Turn Word's document-level 'track changes' setting on or off.

    When enabled, edits made later in Word are recorded as tracked changes.
    """
    filename, error = _prepare(filename)
    if error:
        return error
    try:
        doc = Document(filename)
        R.set_track_changes(doc, enabled)
        doc.save(filename)
        state = "enabled" if enabled else "disabled"
        return f"Track changes {state} for {filename}."
    except Exception as e:
        return f"Failed to toggle track changes: {str(e)}"


async def add_tracked_insertion(
    filename: str,
    text: str,
    author: str = "WaDocx",
    date: Optional[str] = None,
    style: Optional[str] = None,
) -> str:
    """Append a new paragraph whose text is recorded as a tracked insertion."""
    filename, error = _prepare(filename)
    if error:
        return error
    if not text:
        return "Invalid parameter: text must not be empty."
    try:
        doc = Document(filename)
        R.add_tracked_insertion_paragraph(doc, text, author, date or R._now_iso(), style)
        doc.save(filename)
        return f"Tracked insertion added to {filename} (author: {author})."
    except Exception as e:
        return f"Failed to add tracked insertion: {str(e)}"


async def mark_text_as_deleted(
    filename: str,
    search_text: str,
    author: str = "WaDocx",
    date: Optional[str] = None,
) -> str:
    """Mark every occurrence of ``search_text`` as a tracked deletion."""
    filename, error = _prepare(filename)
    if error:
        return error
    if not search_text:
        return "Invalid parameter: search_text must not be empty."
    try:
        doc = Document(filename)
        count = R.mark_paragraph_text_deleted(doc, search_text, author, date or R._now_iso())
        if not count:
            return f"Text '{search_text}' not found in {filename}."
        doc.save(filename)
        return (
            f"Marked {count} occurrence(s) of '{search_text}' as tracked "
            f"deletion(s) in {filename} (author: {author})."
        )
    except Exception as e:
        return f"Failed to mark text deleted: {str(e)}"


async def tracked_search_and_replace(
    filename: str,
    find_text: str,
    replace_text: str,
    author: str = "WaDocx",
    date: Optional[str] = None,
) -> str:
    """Replace text as a tracked change (deletion of old + insertion of new)."""
    filename, error = _prepare(filename)
    if error:
        return error
    if not find_text:
        return "Invalid parameter: find_text must not be empty."
    try:
        doc = Document(filename)
        count = R.tracked_replace(doc, find_text, replace_text, author, date or R._now_iso())
        if not count:
            return f"Text '{find_text}' not found in {filename}."
        doc.save(filename)
        return (
            f"Tracked-replaced {count} occurrence(s) of '{find_text}' with "
            f"'{replace_text}' in {filename} (author: {author})."
        )
    except Exception as e:
        return f"Failed to perform tracked replace: {str(e)}"


async def accept_all_changes(filename: str) -> str:
    """Accept every tracked change in the document (keep insertions, drop deletions)."""
    filename, error = _prepare(filename)
    if error:
        return error
    try:
        doc = Document(filename)
        resolved = R.accept_all_changes(doc)
        doc.save(filename)
        return f"Accepted {resolved} tracked change element(s) in {filename}."
    except Exception as e:
        return f"Failed to accept changes: {str(e)}"


async def reject_all_changes(filename: str) -> str:
    """Reject every tracked change in the document (drop insertions, restore deletions)."""
    filename, error = _prepare(filename)
    if error:
        return error
    try:
        doc = Document(filename)
        resolved = R.reject_all_changes(doc)
        doc.save(filename)
        return f"Rejected {resolved} tracked change element(s) in {filename}."
    except Exception as e:
        return f"Failed to reject changes: {str(e)}"


async def get_revision_counts(filename: str) -> str:
    """Return the number of tracked insertions and deletions in the document."""
    import json

    filename = ensure_docx_extension(filename)
    if not os.path.exists(filename):
        return f"Document {filename} does not exist"
    try:
        doc = Document(filename)
        counts = R.count_revisions(doc)
        counts["filename"] = filename
        return json.dumps(counts, indent=2)
    except Exception as e:
        return f"Failed to count revisions: {str(e)}"
