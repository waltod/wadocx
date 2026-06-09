"""
Content tools for WaDocx MCP.

These tools add various types of content to Word documents,
including headings, paragraphs, tables, images, and page breaks.
"""
import os
from typing import List, Optional, Dict, Any
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from word_document_server.utils.file_utils import check_file_writeable, ensure_docx_extension
from word_document_server.utils.document_utils import find_and_replace_text, insert_header_near_text, insert_numbered_list_near_text, insert_line_or_paragraph_near_text, replace_paragraph_block_below_header, replace_block_between_manual_anchors, build_toc_instruction, normalize_hex_color
from word_document_server.core.styles import ensure_heading_style, ensure_table_style


def _apply_run_formatting(
    run,
    font_name: Optional[str] = None,
    font_size: Optional[int] = None,
    bold: Optional[bool] = None,
    italic: Optional[bool] = None,
    color: Optional[str] = None,
) -> None:
    """Apply common font formatting to a run."""
    if font_name:
        run.font.name = font_name
    if font_size:
        run.font.size = Pt(font_size)
    if bold is not None:
        run.font.bold = bold
    if italic is not None:
        run.font.italic = italic
    normalized = normalize_hex_color(color)
    if normalized:
        run.font.color.rgb = RGBColor.from_string(normalized)


def _set_paragraph_alignment(paragraph, alignment: Optional[str]) -> None:
    """Apply paragraph alignment if requested."""
    if not alignment:
        return

    alignment_map = {
        "left": WD_ALIGN_PARAGRAPH.LEFT,
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "right": WD_ALIGN_PARAGRAPH.RIGHT,
        "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
    }
    alignment_value = alignment_map.get(alignment.lower())
    if alignment_value is not None:
        paragraph.alignment = alignment_value


def _validate_section_index(doc, section_index: int) -> Optional[str]:
    """Validate section index against document sections."""
    if section_index < 0 or section_index >= len(doc.sections):
        return f"Invalid section index. Document has {len(doc.sections)} section(s) (0-{len(doc.sections)-1})."
    return None


def _ensure_update_fields_on_open(doc) -> None:
    """Ask Word to update fields when the document opens."""
    settings_element = doc.settings.element
    update_fields = settings_element.find(qn("w:updateFields"))
    if update_fields is None:
        update_fields = OxmlElement("w:updateFields")
        settings_element.append(update_fields)
    update_fields.set(qn("w:val"), "true")


def _set_run_properties(run_element, font_name=None, font_size=None, bold=None, italic=None, color=None) -> None:
    """Attach basic run properties directly to an XML run element."""
    has_properties = any(
        value is not None and value != ""
        for value in [font_name, font_size, bold, italic, color]
    )
    if not has_properties:
        return

    rpr = OxmlElement("w:rPr")
    if font_name:
        rfonts = OxmlElement("w:rFonts")
        for attr in ("w:ascii", "w:hAnsi", "w:cs"):
            rfonts.set(qn(attr), font_name)
        rpr.append(rfonts)
    if font_size:
        half_points = str(int(font_size) * 2)
        size_el = OxmlElement("w:sz")
        size_el.set(qn("w:val"), half_points)
        rpr.append(size_el)
        size_cs = OxmlElement("w:szCs")
        size_cs.set(qn("w:val"), half_points)
        rpr.append(size_cs)
    if bold is True:
        rpr.append(OxmlElement("w:b"))
    if italic is True:
        rpr.append(OxmlElement("w:i"))
    normalized_color = normalize_hex_color(color)
    if normalized_color:
        color_el = OxmlElement("w:color")
        color_el.set(qn("w:val"), normalized_color)
        rpr.append(color_el)
    run_element.append(rpr)


def _append_field_code_run(
    paragraph,
    instruction: str,
    display_text: str = "",
    font_name: Optional[str] = None,
    font_size: Optional[int] = None,
    bold: Optional[bool] = None,
    italic: Optional[bool] = None,
    color: Optional[str] = None,
) -> None:
    """Append a native Word field code to a paragraph."""
    p = paragraph._p

    begin_run = OxmlElement("w:r")
    _set_run_properties(begin_run, font_name, font_size, bold, italic, color)
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    fld_char_begin.set(qn("w:dirty"), "true")
    begin_run.append(fld_char_begin)
    p.append(begin_run)

    instr_run = OxmlElement("w:r")
    _set_run_properties(instr_run, font_name, font_size, bold, italic, color)
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = f" {instruction} "
    instr_run.append(instr_text)
    p.append(instr_run)

    separate_run = OxmlElement("w:r")
    fld_char_sep = OxmlElement("w:fldChar")
    fld_char_sep.set(qn("w:fldCharType"), "separate")
    separate_run.append(fld_char_sep)
    p.append(separate_run)

    if display_text:
        result_run = paragraph.add_run(display_text)
        _apply_run_formatting(result_run, font_name, font_size, bold, italic, color)

    end_run = OxmlElement("w:r")
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    end_run.append(fld_char_end)
    p.append(end_run)


def _insert_paragraph_before(paragraph, reference_element) -> None:
    """Move a paragraph before a given body element."""
    reference_element.addprevious(paragraph._element)


def _collect_toc_headings(doc, max_level: int):
    """Return [(paragraph, level, text)] for heading paragraphs up to max_level."""
    headings = []
    for paragraph in doc.paragraphs:
        style = paragraph.style
        name = style.name if style else ""
        if not name or not name.startswith("Heading "):
            continue
        try:
            level = int(name.split(" ")[1])
        except (ValueError, IndexError):
            continue
        text = paragraph.text.strip()
        if level <= max_level and text:
            headings.append((paragraph, level, text))
    return headings


def _ensure_heading_bookmark(paragraph, name: str, bookmark_id: int) -> None:
    """Wrap a paragraph's content in a Word bookmark (idempotent by name)."""
    p = paragraph._p
    # Skip if a bookmark with this name already exists on the paragraph.
    for existing in p.findall(qn("w:bookmarkStart")):
        if existing.get(qn("w:name")) == name:
            return
    start = OxmlElement("w:bookmarkStart")
    start.set(qn("w:id"), str(bookmark_id))
    start.set(qn("w:name"), name)
    end = OxmlElement("w:bookmarkEnd")
    end.set(qn("w:id"), str(bookmark_id))
    # Insert start after the paragraph properties (if any), end at the very end.
    ppr = p.find(qn("w:pPr"))
    if ppr is not None:
        ppr.addnext(start)
    else:
        p.insert(0, start)
    p.append(end)


def _pageref_field_elements(bookmark_name: str):
    """Build the run elements for a ``PAGEREF`` field referencing a bookmark."""
    elements = []

    begin = OxmlElement("w:r")
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    begin.append(fld_begin)
    elements.append(begin)

    instr_run = OxmlElement("w:r")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = f" PAGEREF {bookmark_name} \\h "
    instr_run.append(instr)
    elements.append(instr_run)

    sep = OxmlElement("w:r")
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    sep.append(fld_sep)
    elements.append(sep)

    num_run = OxmlElement("w:r")
    num_t = OxmlElement("w:t")
    num_t.text = "1"
    num_run.append(num_t)
    elements.append(num_run)

    end = OxmlElement("w:r")
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    end.append(fld_end)
    elements.append(end)

    return elements


def _build_toc_entry_paragraph(doc, text: str, level: int, bookmark_name: str):
    """Create a populated TOC entry paragraph (hyperlink + tab + PAGEREF)."""
    para = doc.add_paragraph()
    style_applied = False
    try:
        para.style = f"TOC {level}"
        style_applied = True
    except KeyError:
        pass

    if not style_applied:
        # Fallback for documents whose template lacks built-in "TOC n" styles:
        # indent by level and add a right-aligned tab stop with a dotted leader
        # so the page number sits flush-right with classic TOC dot leaders.
        fmt = para.paragraph_format
        fmt.left_indent = Inches(0.25 * (level - 1))
        try:
            fmt.tab_stops.add_tab_stop(
                Inches(6.0), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS
            )
        except Exception:
            pass

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("w:anchor"), bookmark_name)
    hyperlink.set(qn("w:history"), "1")

    text_run = OxmlElement("w:r")
    text_el = OxmlElement("w:t")
    text_el.set(qn("xml:space"), "preserve")
    text_el.text = text
    text_run.append(text_el)
    hyperlink.append(text_run)

    tab_run = OxmlElement("w:r")
    tab_run.append(OxmlElement("w:tab"))
    hyperlink.append(tab_run)

    for element in _pageref_field_elements(bookmark_name):
        hyperlink.append(element)

    para._p.append(hyperlink)
    return para


def _prepend_field_open_runs(paragraph, instruction: str) -> None:
    """Insert TOC field-begin/instr/separate runs at the start of a paragraph."""
    p = paragraph._p
    anchor = p.find(qn("w:pPr"))

    sep_run = OxmlElement("w:r")
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    sep_run.append(fld_sep)

    instr_run = OxmlElement("w:r")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = f" {instruction} "
    instr_run.append(instr)

    begin_run = OxmlElement("w:r")
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    fld_begin.set(qn("w:dirty"), "true")
    begin_run.append(fld_begin)

    # Insert in order: begin, instr, separate (each after pPr / at index 0).
    if anchor is not None:
        anchor.addnext(sep_run)
        anchor.addnext(instr_run)
        anchor.addnext(begin_run)
    else:
        p.insert(0, sep_run)
        p.insert(0, instr_run)
        p.insert(0, begin_run)


def _append_field_end_run(paragraph) -> None:
    """Append a field-end run to a paragraph (closes the TOC field)."""
    end_run = OxmlElement("w:r")
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    end_run.append(fld_end)
    paragraph._p.append(end_run)


def _build_prerendered_toc(doc, headings, instruction: str):
    """Build a fully populated TOC whose field result contains live entries.

    Each heading is bookmarked and represented by a TOC entry paragraph that
    Word will refresh, but which already renders correctly in viewers that do
    not execute field updates. Returns the list of created entry paragraphs in
    document order (first carries the field-begin runs, last the field-end).
    """
    base_id = _get_bookmark_id(doc)
    entry_paras = []
    for offset, (heading_para, level, text) in enumerate(headings):
        bookmark_name = f"_Toc_wadocx_{base_id + offset}"
        _ensure_heading_bookmark(heading_para, bookmark_name, base_id + offset)
        entry_paras.append(_build_toc_entry_paragraph(doc, text, level, bookmark_name))

    _prepend_field_open_runs(entry_paras[0], instruction)
    _append_field_end_run(entry_paras[-1])
    return entry_paras


def _get_bookmark_id(doc) -> int:
    """Return the next bookmark id for the document."""
    bookmark_ids = []
    for element in doc.element.body.iter():
        if element.tag == qn("w:bookmarkStart"):
            try:
                bookmark_ids.append(int(element.get(qn("w:id"), "0")))
            except ValueError:
                pass
    return (max(bookmark_ids) + 1) if bookmark_ids else 1


def _get_story_container(doc, section_index: int, story_kind: str, story_type: str):
    """Resolve a header/footer container by section and type."""
    section = doc.sections[section_index]
    story_type_normalized = (story_type or "default").lower()

    if story_kind == "header":
        if story_type_normalized == "first":
            section.different_first_page_header_footer = True
            container = section.first_page_header
        elif story_type_normalized == "even":
            doc.settings.odd_and_even_pages_header_footer = True
            container = section.even_page_header
        else:
            container = section.header
    else:
        if story_type_normalized == "first":
            section.different_first_page_header_footer = True
            container = section.first_page_footer
        elif story_type_normalized == "even":
            doc.settings.odd_and_even_pages_header_footer = True
            container = section.even_page_footer
        else:
            container = section.footer

    container.is_linked_to_previous = False
    return container, story_type_normalized


def _clear_story_container(container) -> None:
    """Remove existing paragraphs and tables from a header/footer container."""
    story_element = container._element
    removable_tags = {"p", "tbl"}
    for child in list(story_element.iterchildren()):
        if child.tag.split("}")[-1] in removable_tags:
            story_element.remove(child)


def _set_story_text(
    doc,
    story_kind: str,
    filename: str,
    text: str,
    section_index: int = 0,
    story_type: str = "default",
    font_name: Optional[str] = None,
    font_size: Optional[int] = None,
    bold: Optional[bool] = None,
    italic: Optional[bool] = None,
    color: Optional[str] = None,
    alignment: Optional[str] = None,
) -> str:
    """Create or replace header/footer text for a section."""
    validation_error = _validate_section_index(doc, section_index)
    if validation_error:
        return validation_error

    if story_type.lower() not in {"default", "first", "even"}:
        return "Invalid story type. Valid options: default, first, even."

    if alignment and alignment.lower() not in {"left", "center", "right", "justify"}:
        return "Invalid alignment. Valid options: left, center, right, justify."

    container, story_type_normalized = _get_story_container(doc, section_index, story_kind, story_type)
    _clear_story_container(container)
    paragraph = container.add_paragraph(text)
    _set_paragraph_alignment(paragraph, alignment)

    if not paragraph.runs:
        paragraph.add_run(text)
    for run in paragraph.runs:
        _apply_run_formatting(run, font_name, font_size, bold, italic, color)

    doc.save(filename)
    return (
        f"{story_kind.capitalize()} text updated for section {section_index} "
        f"({story_type_normalized}) in {filename}."
    )


def _get_story_text(doc, section_index: int, story_kind: str, story_type: str) -> str:
    """Return header/footer text for a section."""
    validation_error = _validate_section_index(doc, section_index)
    if validation_error:
        return validation_error

    if story_type.lower() not in {"default", "first", "even"}:
        return "Invalid story type. Valid options: default, first, even."

    container, story_type_normalized = _get_story_container(doc, section_index, story_kind, story_type)
    parts = [paragraph.text for paragraph in container.paragraphs if paragraph.text]
    if not parts:
        return f"No {story_kind} text found for section {section_index} ({story_type_normalized})."
    return "\n".join(parts)


async def add_heading(filename: str, text: str, level: int = 1,
                      font_name: Optional[str] = None, font_size: Optional[int] = None,
                      bold: Optional[bool] = None, italic: Optional[bool] = None,
                      border_bottom: bool = False) -> str:
    """Add a heading to a Word document with optional formatting.

    Args:
        filename: Path to the Word document
        text: Heading text
        level: Heading level (1-9, where 1 is the highest level)
        font_name: Font family (e.g., 'Helvetica')
        font_size: Font size in points (e.g., 14)
        bold: True/False for bold text
        italic: True/False for italic text
        border_bottom: True to add bottom border (for section headers)
    """
    filename = ensure_docx_extension(filename)

    # Ensure level is converted to integer
    try:
        level = int(level)
    except (ValueError, TypeError):
        return "Invalid parameter: level must be an integer between 1 and 9"

    # Validate level range
    if level < 1 or level > 9:
        return f"Invalid heading level: {level}. Level must be between 1 and 9."

    if not os.path.exists(filename):
        return f"Document {filename} does not exist"

    # Check if file is writeable
    is_writeable, error_message = check_file_writeable(filename)
    if not is_writeable:
        # Suggest creating a copy
        return f"Cannot modify document: {error_message}. Consider creating a copy first or creating a new document."

    try:
        doc = Document(filename)

        # Ensure heading styles exist
        ensure_heading_style(doc)

        # Try to add heading with style
        try:
            heading = doc.add_heading(text, level=level)
        except Exception as style_error:
            # If style-based approach fails, use direct formatting
            heading = doc.add_paragraph(text)
            heading.style = doc.styles['Normal']
            if heading.runs:
                run = heading.runs[0]
                run.bold = True
                # Adjust size based on heading level
                if level == 1:
                    run.font.size = Pt(16)
                elif level == 2:
                    run.font.size = Pt(14)
                else:
                    run.font.size = Pt(12)

        # Apply formatting to all runs in the heading
        if any([font_name, font_size, bold is not None, italic is not None]):
            for run in heading.runs:
                if font_name:
                    run.font.name = font_name
                if font_size:
                    run.font.size = Pt(font_size)
                if bold is not None:
                    run.font.bold = bold
                if italic is not None:
                    run.font.italic = italic

        # Add bottom border if requested
        if border_bottom:
            from docx.oxml import OxmlElement
            from docx.oxml.ns import qn

            pPr = heading._element.get_or_add_pPr()
            pBdr = OxmlElement('w:pBdr')

            bottom = OxmlElement('w:bottom')
            bottom.set(qn('w:val'), 'single')
            bottom.set(qn('w:sz'), '4')  # 0.5pt border
            bottom.set(qn('w:space'), '0')
            bottom.set(qn('w:color'), '000000')

            pBdr.append(bottom)
            pPr.append(pBdr)

        doc.save(filename)
        return f"Heading '{text}' (level {level}) added to {filename}"
    except Exception as e:
        return f"Failed to add heading: {str(e)}"


async def add_paragraph(filename: str, text: str, style: Optional[str] = None,
                        font_name: Optional[str] = None, font_size: Optional[int] = None,
                        bold: Optional[bool] = None, italic: Optional[bool] = None,
                        color: Optional[str] = None) -> str:
    """Add a paragraph to a Word document with optional formatting.

    Args:
        filename: Path to the Word document
        text: Paragraph text
        style: Optional paragraph style name
        font_name: Font family (e.g., 'Helvetica', 'Times New Roman')
        font_size: Font size in points (e.g., 14, 36)
        bold: True/False for bold text
        italic: True/False for italic text
        color: RGB color as hex string (e.g., '000000' for black)
    """
    filename = ensure_docx_extension(filename)

    if not os.path.exists(filename):
        return f"Document {filename} does not exist"

    # Check if file is writeable
    is_writeable, error_message = check_file_writeable(filename)
    if not is_writeable:
        # Suggest creating a copy
        return f"Cannot modify document: {error_message}. Consider creating a copy first or creating a new document."

    try:
        doc = Document(filename)
        paragraph = doc.add_paragraph(text)

        if style:
            try:
                paragraph.style = style
            except KeyError:
                # Style doesn't exist, use normal and report it
                paragraph.style = doc.styles['Normal']
                doc.save(filename)
                return f"Style '{style}' not found, paragraph added with default style to {filename}"

        # Apply formatting to all runs in the paragraph
        if any([font_name, font_size, bold is not None, italic is not None, color]):
            for run in paragraph.runs:
                if font_name:
                    run.font.name = font_name
                if font_size:
                    run.font.size = Pt(font_size)
                if bold is not None:
                    run.font.bold = bold
                if italic is not None:
                    run.font.italic = italic
                if color:
                    # Remove any '#' prefix if present
                    color_hex = color.lstrip('#')
                    run.font.color.rgb = RGBColor.from_string(color_hex)

        doc.save(filename)
        return f"Paragraph added to {filename}"
    except Exception as e:
        return f"Failed to add paragraph: {str(e)}"


async def add_table(filename: str, rows: int, cols: int, data: Optional[List[List[str]]] = None) -> str:
    """Add a table to a Word document.
    
    Args:
        filename: Path to the Word document
        rows: Number of rows in the table
        cols: Number of columns in the table
        data: Optional 2D array of data to fill the table
    """
    filename = ensure_docx_extension(filename)
    
    if not os.path.exists(filename):
        return f"Document {filename} does not exist"
    
    # Check if file is writeable
    is_writeable, error_message = check_file_writeable(filename)
    if not is_writeable:
        # Suggest creating a copy
        return f"Cannot modify document: {error_message}. Consider creating a copy first or creating a new document."
    
    try:
        doc = Document(filename)
        table = doc.add_table(rows=rows, cols=cols)
        
        # Try to set the table style
        try:
            table.style = 'Table Grid'
        except KeyError:
            # If style doesn't exist, add basic borders
            pass
        
        # Fill table with data if provided
        if data:
            for i, row_data in enumerate(data):
                if i >= rows:
                    break
                for j, cell_text in enumerate(row_data):
                    if j >= cols:
                        break
                    table.cell(i, j).text = str(cell_text)
        
        doc.save(filename)
        return f"Table ({rows}x{cols}) added to {filename}"
    except Exception as e:
        return f"Failed to add table: {str(e)}"


async def add_picture(filename: str, image_path: str, width: Optional[float] = None) -> str:
    """Add an image to a Word document.
    
    Args:
        filename: Path to the Word document
        image_path: Path to the image file
        width: Optional width in inches (proportional scaling)
    """
    filename = ensure_docx_extension(filename)
    
    # Validate document existence
    if not os.path.exists(filename):
        return f"Document {filename} does not exist"
    
    # Get absolute paths for better diagnostics
    abs_filename = os.path.abspath(filename)
    abs_image_path = os.path.abspath(image_path)
    
    # Validate image existence with improved error message
    if not os.path.exists(abs_image_path):
        return f"Image file not found: {abs_image_path}"
    
    # Check image file size
    try:
        image_size = os.path.getsize(abs_image_path) / 1024  # Size in KB
        if image_size <= 0:
            return f"Image file appears to be empty: {abs_image_path} (0 KB)"
    except Exception as size_error:
        return f"Error checking image file: {str(size_error)}"
    
    # Check if file is writeable
    is_writeable, error_message = check_file_writeable(abs_filename)
    if not is_writeable:
        return f"Cannot modify document: {error_message}. Consider creating a copy first or creating a new document."
    
    try:
        doc = Document(abs_filename)
        # Additional diagnostic info
        diagnostic = f"Attempting to add image ({abs_image_path}, {image_size:.2f} KB) to document ({abs_filename})"
        
        try:
            if width:
                doc.add_picture(abs_image_path, width=Inches(width))
            else:
                doc.add_picture(abs_image_path)
            doc.save(abs_filename)
            return f"Picture {image_path} added to {filename}"
        except Exception as inner_error:
            # More detailed error for the specific operation
            error_type = type(inner_error).__name__
            error_msg = str(inner_error)
            return f"Failed to add picture: {error_type} - {error_msg or 'No error details available'}\nDiagnostic info: {diagnostic}"
    except Exception as outer_error:
        # Fallback error handling
        error_type = type(outer_error).__name__
        error_msg = str(outer_error)
        return f"Document processing error: {error_type} - {error_msg or 'No error details available'}"


async def add_page_break(filename: str) -> str:
    """Add a page break to the document.
    
    Args:
        filename: Path to the Word document
    """
    filename = ensure_docx_extension(filename)
    
    if not os.path.exists(filename):
        return f"Document {filename} does not exist"
    
    # Check if file is writeable
    is_writeable, error_message = check_file_writeable(filename)
    if not is_writeable:
        return f"Cannot modify document: {error_message}. Consider creating a copy first."
    
    try:
        doc = Document(filename)
        doc.add_page_break()
        doc.save(filename)
        return f"Page break added to {filename}."
    except Exception as e:
        return f"Failed to add page break: {str(e)}"


async def set_document_header(
    filename: str,
    text: str,
    section_index: int = 0,
    header_type: str = "default",
    font_name: Optional[str] = None,
    font_size: Optional[int] = None,
    bold: Optional[bool] = None,
    italic: Optional[bool] = None,
    color: Optional[str] = None,
    alignment: Optional[str] = None,
) -> str:
    """Create or replace page header text for a document section."""
    filename = ensure_docx_extension(filename)

    if not os.path.exists(filename):
        return f"Document {filename} does not exist"

    is_writeable, error_message = check_file_writeable(filename)
    if not is_writeable:
        return f"Cannot modify document: {error_message}. Consider creating a copy first."

    try:
        doc = Document(filename)
        return _set_story_text(
            doc,
            "header",
            filename,
            text,
            int(section_index),
            header_type,
            font_name,
            int(font_size) if font_size is not None else None,
            bold,
            italic,
            color,
            alignment,
        )
    except Exception as e:
        return f"Failed to set header text: {str(e)}"


async def get_document_header(filename: str, section_index: int = 0, header_type: str = "default") -> str:
    """Read page header text for a document section."""
    filename = ensure_docx_extension(filename)

    if not os.path.exists(filename):
        return f"Document {filename} does not exist"

    try:
        doc = Document(filename)
        return _get_story_text(doc, int(section_index), "header", header_type)
    except Exception as e:
        return f"Failed to get header text: {str(e)}"


async def set_document_footer(
    filename: str,
    text: str,
    section_index: int = 0,
    footer_type: str = "default",
    font_name: Optional[str] = None,
    font_size: Optional[int] = None,
    bold: Optional[bool] = None,
    italic: Optional[bool] = None,
    color: Optional[str] = None,
    alignment: Optional[str] = None,
) -> str:
    """Create or replace page footer text for a document section."""
    filename = ensure_docx_extension(filename)

    if not os.path.exists(filename):
        return f"Document {filename} does not exist"

    is_writeable, error_message = check_file_writeable(filename)
    if not is_writeable:
        return f"Cannot modify document: {error_message}. Consider creating a copy first."

    try:
        doc = Document(filename)
        return _set_story_text(
            doc,
            "footer",
            filename,
            text,
            int(section_index),
            footer_type,
            font_name,
            int(font_size) if font_size is not None else None,
            bold,
            italic,
            color,
            alignment,
        )
    except Exception as e:
        return f"Failed to set footer text: {str(e)}"


async def get_document_footer(filename: str, section_index: int = 0, footer_type: str = "default") -> str:
    """Read page footer text for a document section."""
    filename = ensure_docx_extension(filename)

    if not os.path.exists(filename):
        return f"Document {filename} does not exist"

    try:
        doc = Document(filename)
        return _get_story_text(doc, int(section_index), "footer", footer_type)
    except Exception as e:
        return f"Failed to get footer text: {str(e)}"


async def add_live_table_of_contents(
    filename: str,
    title: str = "Contents",
    max_level: int = 3,
    insert_at_start: bool = True,
    add_page_break_after: bool = False,
    toc_style: str = "dotted",
    prerender: bool = True,
) -> str:
    """Insert a native Word TOC field that can be refreshed in Word.

    When ``prerender`` is True (default) and the document already contains
    headings, the TOC field result is populated with live entries (hyperlink +
    tab + PAGEREF page number) so the TOC is visible immediately — including in
    viewers that do not execute field updates (LibreOffice/headless PDF export,
    Google Docs preview). Word still refreshes the page numbers on open.
    """
    filename = ensure_docx_extension(filename)

    if not os.path.exists(filename):
        return f"Document {filename} does not exist"

    is_writeable, error_message = check_file_writeable(filename)
    if not is_writeable:
        return f"Cannot modify document: {error_message}. Consider creating a copy first."

    try:
        max_level = max(1, min(int(max_level), 9))
    except (ValueError, TypeError):
        return "Invalid parameter: max_level must be an integer between 1 and 9"

    try:
        doc = Document(filename)
        _ensure_update_fields_on_open(doc)
        instruction = build_toc_instruction(max_level=max_level, toc_style=toc_style)

        toc_title_para = doc.add_paragraph(title) if title else None
        if toc_title_para is not None and title:
            try:
                toc_title_para.style = "TOC Heading"
            except KeyError:
                if toc_title_para.runs:
                    toc_title_para.runs[0].bold = True

        headings = _collect_toc_headings(doc, max_level) if prerender else []
        entry_count = len(headings)

        if headings:
            # Populated TOC: one paragraph per heading, wrapped as a field.
            toc_paras = _build_prerendered_toc(doc, headings, instruction)
        else:
            # Empty/placeholder TOC field (no headings yet, or prerender disabled).
            toc_para = doc.add_paragraph()
            _append_field_code_run(
                toc_para,
                instruction,
                display_text="Right-click to update field.",
            )
            toc_paras = [toc_para]

        if add_page_break_after:
            page_break_para = doc.add_paragraph()
            page_break_para.add_run().add_break(WD_BREAK.PAGE)
        else:
            page_break_para = None

        if insert_at_start and doc.paragraphs:
            first_element = doc.paragraphs[0]._element
            ordered = [item for item in [toc_title_para, *toc_paras, page_break_para] if item is not None]
            # addprevious keeps insertion order, so iterate in visual order.
            for para in ordered:
                _insert_paragraph_before(para, first_element)

        doc.save(filename)
        if entry_count:
            return (
                f"Live table of contents with {entry_count} entr"
                f"{'y' if entry_count == 1 else 'ies'} inserted into {filename} "
                f"({toc_style}, pre-rendered)."
            )
        return (
            f"Live table of contents field inserted into {filename} ({toc_style}). "
            f"No headings found yet — it will populate when you add Heading-styled "
            f"paragraphs and refresh the field in Word."
        )
    except Exception as e:
        return f"Failed to insert live table of contents: {str(e)}"


async def set_document_header_page_number(
    filename: str,
    prefix_text: str = "",
    suffix_text: str = "",
    section_index: int = 0,
    header_type: str = "default",
    alignment: str = "right",
    font_name: Optional[str] = None,
    font_size: Optional[int] = None,
    bold: Optional[bool] = None,
    italic: Optional[bool] = None,
    color: Optional[str] = None,
) -> str:
    """Create or replace a page header containing a live PAGE field."""
    filename = ensure_docx_extension(filename)

    if not os.path.exists(filename):
        return f"Document {filename} does not exist"

    is_writeable, error_message = check_file_writeable(filename)
    if not is_writeable:
        return f"Cannot modify document: {error_message}. Consider creating a copy first."

    try:
        doc = Document(filename)
        validation_error = _validate_section_index(doc, int(section_index))
        if validation_error:
            return validation_error

        container, story_type_normalized = _get_story_container(doc, int(section_index), "header", header_type)
        _clear_story_container(container)
        paragraph = container.add_paragraph()
        _set_paragraph_alignment(paragraph, alignment)
        if prefix_text:
            prefix_run = paragraph.add_run(prefix_text)
            _apply_run_formatting(prefix_run, font_name, font_size, bold, italic, color)
        _ensure_update_fields_on_open(doc)
        _append_field_code_run(
            paragraph,
            "PAGE",
            display_text="1",
            font_name=font_name,
            font_size=int(font_size) if font_size is not None else None,
            bold=bold,
            italic=italic,
            color=color,
        )
        if suffix_text:
            suffix_run = paragraph.add_run(suffix_text)
            _apply_run_formatting(suffix_run, font_name, font_size, bold, italic, color)
        doc.save(filename)
        return f"Header page number field updated for section {section_index} ({story_type_normalized}) in {filename}."
    except Exception as e:
        return f"Failed to set header page number field: {str(e)}"


async def set_document_footer_page_number(
    filename: str,
    prefix_text: str = "",
    suffix_text: str = "",
    section_index: int = 0,
    footer_type: str = "default",
    alignment: str = "right",
    font_name: Optional[str] = None,
    font_size: Optional[int] = None,
    bold: Optional[bool] = None,
    italic: Optional[bool] = None,
    color: Optional[str] = None,
) -> str:
    """Create or replace a page footer containing a live PAGE field."""
    filename = ensure_docx_extension(filename)

    if not os.path.exists(filename):
        return f"Document {filename} does not exist"

    is_writeable, error_message = check_file_writeable(filename)
    if not is_writeable:
        return f"Cannot modify document: {error_message}. Consider creating a copy first."

    try:
        doc = Document(filename)
        validation_error = _validate_section_index(doc, int(section_index))
        if validation_error:
            return validation_error

        container, story_type_normalized = _get_story_container(doc, int(section_index), "footer", footer_type)
        _clear_story_container(container)
        paragraph = container.add_paragraph()
        _set_paragraph_alignment(paragraph, alignment)
        if prefix_text:
            prefix_run = paragraph.add_run(prefix_text)
            _apply_run_formatting(prefix_run, font_name, font_size, bold, italic, color)
        _ensure_update_fields_on_open(doc)
        _append_field_code_run(
            paragraph,
            "PAGE",
            display_text="1",
            font_name=font_name,
            font_size=int(font_size) if font_size is not None else None,
            bold=bold,
            italic=italic,
            color=color,
        )
        if suffix_text:
            suffix_run = paragraph.add_run(suffix_text)
            _apply_run_formatting(suffix_run, font_name, font_size, bold, italic, color)
        doc.save(filename)
        return f"Footer page number field updated for section {section_index} ({story_type_normalized}) in {filename}."
    except Exception as e:
        return f"Failed to set footer page number field: {str(e)}"


async def insert_omml_equation(
    filename: str,
    equation_text: str,
    paragraph_index: Optional[int] = None,
    position: str = "after",
) -> str:
    """Insert a native Word OMML equation paragraph."""
    filename = ensure_docx_extension(filename)

    if not os.path.exists(filename):
        return f"Document {filename} does not exist"

    is_writeable, error_message = check_file_writeable(filename)
    if not is_writeable:
        return f"Cannot modify document: {error_message}. Consider creating a copy first."

    if position not in {"before", "after"}:
        return "Invalid position. Valid options: before, after."

    try:
        doc = Document(filename)
        equation_para = doc.add_paragraph()
        o_math_para = OxmlElement("m:oMathPara")
        o_math = OxmlElement("m:oMath")
        math_run = OxmlElement("m:r")
        word_run = OxmlElement("w:r")
        text_el = OxmlElement("w:t")
        text_el.text = equation_text
        word_run.append(text_el)
        math_run.append(word_run)
        o_math.append(math_run)
        o_math_para.append(o_math)
        equation_para._p.append(o_math_para)

        if paragraph_index is not None:
            paragraph_index = int(paragraph_index)
            if paragraph_index < 0 or paragraph_index >= len(doc.paragraphs):
                return f"Invalid paragraph index. Document has {len(doc.paragraphs)} paragraphs (0-{len(doc.paragraphs)-1})."
            anchor = doc.paragraphs[paragraph_index]
            if position == "before":
                anchor._element.addprevious(equation_para._element)
            else:
                anchor._element.addnext(equation_para._element)

        doc.save(filename)
        return f"OMML equation inserted into {filename}."
    except Exception as e:
        return f"Failed to insert OMML equation: {str(e)}"


async def add_bookmark_to_paragraph(filename: str, paragraph_index: int, bookmark_name: str) -> str:
    """Attach a native Word bookmark to a paragraph."""
    filename = ensure_docx_extension(filename)

    if not os.path.exists(filename):
        return f"Document {filename} does not exist"

    is_writeable, error_message = check_file_writeable(filename)
    if not is_writeable:
        return f"Cannot modify document: {error_message}. Consider creating a copy first."

    try:
        doc = Document(filename)
        paragraph_index = int(paragraph_index)
        if paragraph_index < 0 or paragraph_index >= len(doc.paragraphs):
            return f"Invalid paragraph index. Document has {len(doc.paragraphs)} paragraphs (0-{len(doc.paragraphs)-1})."

        paragraph = doc.paragraphs[paragraph_index]
        bookmark_id = str(_get_bookmark_id(doc))

        start = OxmlElement("w:bookmarkStart")
        start.set(qn("w:id"), bookmark_id)
        start.set(qn("w:name"), bookmark_name)
        paragraph._p.insert(0, start)

        end = OxmlElement("w:bookmarkEnd")
        end.set(qn("w:id"), bookmark_id)
        paragraph._p.append(end)

        doc.save(filename)
        return f"Bookmark '{bookmark_name}' added to paragraph {paragraph_index} in {filename}."
    except Exception as e:
        return f"Failed to add bookmark: {str(e)}"


async def add_internal_hyperlink(
    filename: str,
    paragraph_index: int,
    link_text: str,
    bookmark_name: str,
) -> str:
    """Append an internal hyperlink pointing to a bookmark."""
    filename = ensure_docx_extension(filename)

    if not os.path.exists(filename):
        return f"Document {filename} does not exist"

    is_writeable, error_message = check_file_writeable(filename)
    if not is_writeable:
        return f"Cannot modify document: {error_message}. Consider creating a copy first."

    try:
        doc = Document(filename)
        paragraph_index = int(paragraph_index)
        if paragraph_index < 0 or paragraph_index >= len(doc.paragraphs):
            return f"Invalid paragraph index. Document has {len(doc.paragraphs)} paragraphs (0-{len(doc.paragraphs)-1})."

        paragraph = doc.paragraphs[paragraph_index]
        hyperlink = OxmlElement("w:hyperlink")
        hyperlink.set(qn("w:anchor"), bookmark_name)

        run = OxmlElement("w:r")
        rpr = OxmlElement("w:rPr")
        rstyle = OxmlElement("w:rStyle")
        rstyle.set(qn("w:val"), "Hyperlink")
        rpr.append(rstyle)
        run.append(rpr)
        text_el = OxmlElement("w:t")
        text_el.text = link_text
        run.append(text_el)
        hyperlink.append(run)
        paragraph._p.append(hyperlink)

        doc.save(filename)
        return f"Internal hyperlink to '{bookmark_name}' added to paragraph {paragraph_index} in {filename}."
    except Exception as e:
        return f"Failed to add internal hyperlink: {str(e)}"


async def add_table_of_contents(filename: str, title: str = "Table of Contents", max_level: int = 3) -> str:
    """Add a table of contents to a Word document based on heading styles.
    
    Args:
        filename: Path to the Word document
        title: Optional title for the table of contents
        max_level: Maximum heading level to include (1-9)
    """
    filename = ensure_docx_extension(filename)
    
    if not os.path.exists(filename):
        return f"Document {filename} does not exist"
    
    # Check if file is writeable
    is_writeable, error_message = check_file_writeable(filename)
    if not is_writeable:
        return f"Cannot modify document: {error_message}. Consider creating a copy first."
    
    # Delegate to the native, non-destructive live-TOC implementation.
    #
    # The previous implementation rebuilt the entire document from scratch by
    # copying only ``paragraph.text`` into a brand-new ``Document()``. That
    # destroyed all run formatting, images, hyperlinks, headers/footers,
    # sections and footnotes, duplicated every heading, and moved all tables to
    # the end of the file. Inserting a genuine Word TOC field in place preserves
    # the document and produces a TOC that Word can refresh.
    return await add_live_table_of_contents(
        filename=filename,
        title=title,
        max_level=max_level,
        insert_at_start=True,
        add_page_break_after=True,
        toc_style="dotted",
    )


async def delete_paragraph(filename: str, paragraph_index: int) -> str:
    """Delete a paragraph from a document.
    
    Args:
        filename: Path to the Word document
        paragraph_index: Index of the paragraph to delete (0-based)
    """
    filename = ensure_docx_extension(filename)
    
    if not os.path.exists(filename):
        return f"Document {filename} does not exist"
    
    # Check if file is writeable
    is_writeable, error_message = check_file_writeable(filename)
    if not is_writeable:
        return f"Cannot modify document: {error_message}. Consider creating a copy first."
    
    try:
        doc = Document(filename)
        
        # Validate paragraph index
        if paragraph_index < 0 or paragraph_index >= len(doc.paragraphs):
            return f"Invalid paragraph index. Document has {len(doc.paragraphs)} paragraphs (0-{len(doc.paragraphs)-1})."
        
        # Delete the paragraph (by removing its content and setting it empty)
        # Note: python-docx doesn't support true paragraph deletion, this is a workaround
        paragraph = doc.paragraphs[paragraph_index]
        p = paragraph._p
        p.getparent().remove(p)
        
        doc.save(filename)
        return f"Paragraph at index {paragraph_index} deleted successfully."
    except Exception as e:
        return f"Failed to delete paragraph: {str(e)}"


async def search_and_replace(filename: str, find_text: str, replace_text: str) -> str:
    """Search for text and replace all occurrences.
    
    Args:
        filename: Path to the Word document
        find_text: Text to search for
        replace_text: Text to replace with
    """
    filename = ensure_docx_extension(filename)
    
    if not os.path.exists(filename):
        return f"Document {filename} does not exist"
    
    # Check if file is writeable
    is_writeable, error_message = check_file_writeable(filename)
    if not is_writeable:
        return f"Cannot modify document: {error_message}. Consider creating a copy first."
    
    try:
        doc = Document(filename)
        
        # Perform find and replace
        count = find_and_replace_text(doc, find_text, replace_text)
        
        if count > 0:
            doc.save(filename)
            return f"Replaced {count} occurrence(s) of '{find_text}' with '{replace_text}'."
        else:
            return f"No occurrences of '{find_text}' found."
    except Exception as e:
        return f"Failed to search and replace: {str(e)}"

async def insert_header_near_text_tool(filename: str, target_text: str = None, header_title: str = "", position: str = 'after', header_style: str = 'Heading 1', target_paragraph_index: int = None) -> str:
    """Insert a header (with specified style) before or after the target paragraph. Specify by text or paragraph index."""
    return insert_header_near_text(filename, target_text, header_title, position, header_style, target_paragraph_index)

async def insert_numbered_list_near_text_tool(filename: str, target_text: str = None, list_items: list = None, position: str = 'after', target_paragraph_index: int = None, bullet_type: str = 'bullet') -> str:
    """Insert a bulleted or numbered list before or after the target paragraph. Specify by text or paragraph index."""
    return insert_numbered_list_near_text(filename, target_text, list_items, position, target_paragraph_index, bullet_type)

async def insert_line_or_paragraph_near_text_tool(filename: str, target_text: str = None, line_text: str = "", position: str = 'after', line_style: str = None, target_paragraph_index: int = None) -> str:
    """Insert a new line or paragraph (with specified or matched style) before or after the target paragraph. Specify by text or paragraph index."""
    return insert_line_or_paragraph_near_text(filename, target_text, line_text, position, line_style, target_paragraph_index)

async def replace_paragraph_block_below_header_tool(filename: str, header_text: str, new_paragraphs: list, detect_block_end_fn=None) -> str:
    """Reemplaza el bloque de párrafos debajo de un encabezado, evitando modificar TOC."""
    return replace_paragraph_block_below_header(filename, header_text, new_paragraphs, detect_block_end_fn)

async def replace_block_between_manual_anchors_tool(filename: str, start_anchor_text: str, new_paragraphs: list, end_anchor_text: str = None, match_fn=None, new_paragraph_style: str = None) -> str:
    """Replace all content between start_anchor_text and end_anchor_text (or next logical header if not provided)."""
    return replace_block_between_manual_anchors(filename, start_anchor_text, new_paragraphs, end_anchor_text, match_fn, new_paragraph_style)



# ---------------------------------------------------------------------------
# New tools: external hyperlinks, statistics, page setup, captions / figures
# ---------------------------------------------------------------------------

async def add_hyperlink(
    filename: str,
    url: str,
    text: Optional[str] = None,
    paragraph_index: Optional[int] = None,
    color: Optional[str] = "0563C1",
    underline: bool = True,
) -> str:
    """Insert a clickable external hyperlink into a Word document.

    Args:
        filename: Path to the Word document.
        url: The destination URL (e.g. https://example.com).
        text: Visible link text (defaults to the URL).
        paragraph_index: Append the link to this existing paragraph. If omitted,
            a new paragraph is added at the end of the document.
        color: Hex color for the link text (default Word hyperlink blue).
        underline: Whether to underline the link text.
    """
    filename = ensure_docx_extension(filename)
    if not os.path.exists(filename):
        return f"Document {filename} does not exist"

    is_writeable, error_message = check_file_writeable(filename)
    if not is_writeable:
        return f"Cannot modify document: {error_message}. Consider creating a copy first."

    if not url or not str(url).strip():
        return "Invalid parameter: url must not be empty."

    try:
        normalized_color = normalize_hex_color(color)
    except ValueError as exc:
        return str(exc)

    try:
        from docx.opc.constants import RELATIONSHIP_TYPE as RT

        doc = Document(filename)
        link_text = text if text else url

        if paragraph_index is not None:
            if paragraph_index < 0 or paragraph_index >= len(doc.paragraphs):
                return (
                    f"Invalid paragraph index. Document has {len(doc.paragraphs)} "
                    f"paragraphs (0-{len(doc.paragraphs) - 1})."
                )
            paragraph = doc.paragraphs[paragraph_index]
        else:
            paragraph = doc.add_paragraph()

        part = paragraph.part
        r_id = part.relate_to(url, RT.HYPERLINK, is_external=True)

        hyperlink = OxmlElement("w:hyperlink")
        hyperlink.set(qn("r:id"), r_id)
        hyperlink.set(qn("w:history"), "1")

        new_run = OxmlElement("w:r")
        rpr = OxmlElement("w:rPr")
        rstyle = OxmlElement("w:rStyle")
        rstyle.set(qn("w:val"), "Hyperlink")
        rpr.append(rstyle)
        if normalized_color:
            color_el = OxmlElement("w:color")
            color_el.set(qn("w:val"), normalized_color)
            rpr.append(color_el)
        if underline:
            u_el = OxmlElement("w:u")
            u_el.set(qn("w:val"), "single")
            rpr.append(u_el)
        new_run.append(rpr)

        text_el = OxmlElement("w:t")
        text_el.set(qn("xml:space"), "preserve")
        text_el.text = link_text
        new_run.append(text_el)
        hyperlink.append(new_run)

        paragraph._p.append(hyperlink)
        doc.save(filename)
        return f"Hyperlink to {url} added to {filename}."
    except Exception as e:
        return f"Failed to add hyperlink: {str(e)}"


async def get_document_statistics(filename: str) -> str:
    """Return word/character/structure statistics for a Word document as JSON.

    Counts words, characters (with and without spaces), paragraphs, headings (by
    level), tables, images, footnotes, hyperlinks, sections, and a rough page
    estimate.
    """
    import json as _json

    filename = ensure_docx_extension(filename)
    if not os.path.exists(filename):
        return f"Document {filename} does not exist"

    try:
        doc = Document(filename)

        words = 0
        chars_no_spaces = 0
        chars_with_spaces = 0
        paragraph_count = 0
        empty_paragraphs = 0
        headings_by_level: Dict[str, int] = {}

        def _tally(text: str) -> None:
            nonlocal words, chars_no_spaces, chars_with_spaces
            words += len(text.split())
            chars_with_spaces += len(text)
            chars_no_spaces += len(text.replace(" ", "").replace("\t", ""))

        for paragraph in doc.paragraphs:
            paragraph_count += 1
            text = paragraph.text
            if not text.strip():
                empty_paragraphs += 1
            _tally(text)
            style_name = paragraph.style.name if paragraph.style else ""
            if style_name.startswith("Heading "):
                headings_by_level[style_name] = headings_by_level.get(style_name, 0) + 1

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        _tally(paragraph.text)

        body_xml = doc.element.body
        image_count = len(body_xml.findall(".//" + qn("w:drawing")))
        hyperlink_count = len(body_xml.findall(".//" + qn("w:hyperlink")))

        footnote_count = 0
        try:
            for rel in doc.part.rels.values():
                if "footnotes" in rel.reltype:
                    fn_part = rel.target_part
                    fn_root = getattr(fn_part, "_element", None)
                    if fn_root is not None:
                        notes = fn_root.findall(qn("w:footnote"))
                        footnote_count = max(0, len(notes) - 2)
                    break
        except Exception:
            footnote_count = 0

        stats = {
            "filename": filename,
            "words": words,
            "characters_with_spaces": chars_with_spaces,
            "characters_no_spaces": chars_no_spaces,
            "paragraphs": paragraph_count,
            "non_empty_paragraphs": paragraph_count - empty_paragraphs,
            "headings": sum(headings_by_level.values()),
            "headings_by_level": headings_by_level,
            "tables": len(doc.tables),
            "images": image_count,
            "hyperlinks": hyperlink_count,
            "footnotes": footnote_count,
            "sections": len(doc.sections),
            "estimated_pages": max(1, round(words / 500)) if words else 0,
        }
        return _json.dumps(stats, indent=2)
    except Exception as e:
        return f"Failed to compute document statistics: {str(e)}"


async def set_page_setup(
    filename: str,
    section_index: Optional[int] = None,
    orientation: Optional[str] = None,
    page_size: Optional[str] = None,
    margin_top: Optional[float] = None,
    margin_bottom: Optional[float] = None,
    margin_left: Optional[float] = None,
    margin_right: Optional[float] = None,
    units: str = "inches",
) -> str:
    """Configure page orientation, size and margins for one or all sections.

    Args:
        filename: Path to the Word document.
        section_index: Section to modify. If omitted, applies to all sections.
        orientation: 'portrait' or 'landscape'.
        page_size: 'letter', 'legal', 'a4', or 'a3'.
        margin_top/bottom/left/right: Margin sizes in ``units``.
        units: 'inches' (default) or 'cm'.
    """
    from docx.enum.section import WD_ORIENT
    from docx.shared import Inches, Cm

    filename = ensure_docx_extension(filename)
    if not os.path.exists(filename):
        return f"Document {filename} does not exist"

    is_writeable, error_message = check_file_writeable(filename)
    if not is_writeable:
        return f"Cannot modify document: {error_message}. Consider creating a copy first."

    unit = (units or "inches").lower()
    if unit not in {"inches", "cm"}:
        return "Invalid units. Use 'inches' or 'cm'."
    length = Inches if unit == "inches" else Cm

    page_sizes = {
        "letter": (8.5, 11.0),
        "legal": (8.5, 14.0),
        "a4": (8.27, 11.69),
        "a3": (11.69, 16.54),
    }

    try:
        doc = Document(filename)

        if section_index is not None:
            if section_index < 0 or section_index >= len(doc.sections):
                return (
                    f"Invalid section index. Document has {len(doc.sections)} "
                    f"section(s) (0-{len(doc.sections) - 1})."
                )
            sections = [doc.sections[section_index]]
        else:
            sections = list(doc.sections)

        orient = (orientation or "").lower()
        if orientation and orient not in {"portrait", "landscape"}:
            return "Invalid orientation. Use 'portrait' or 'landscape'."

        size_key = (page_size or "").lower()
        if page_size and size_key not in page_sizes:
            return "Invalid page_size. Use 'letter', 'legal', 'a4', or 'a3'."

        for section in sections:
            if page_size:
                w_in, h_in = page_sizes[size_key]
                section.page_width = Inches(w_in)
                section.page_height = Inches(h_in)
            if orientation:
                if orient == "landscape":
                    section.orientation = WD_ORIENT.LANDSCAPE
                    if section.page_width < section.page_height:
                        section.page_width, section.page_height = (
                            section.page_height,
                            section.page_width,
                        )
                else:
                    section.orientation = WD_ORIENT.PORTRAIT
                    if section.page_width > section.page_height:
                        section.page_width, section.page_height = (
                            section.page_height,
                            section.page_width,
                        )
            if margin_top is not None:
                section.top_margin = length(margin_top)
            if margin_bottom is not None:
                section.bottom_margin = length(margin_bottom)
            if margin_left is not None:
                section.left_margin = length(margin_left)
            if margin_right is not None:
                section.right_margin = length(margin_right)

        doc.save(filename)
        scope = (
            f"section {section_index}" if section_index is not None
            else f"all {len(sections)} section(s)"
        )
        return f"Page setup updated for {scope} in {filename}."
    except Exception as e:
        return f"Failed to update page setup: {str(e)}"


def _append_seq_caption_runs(paragraph, label: str) -> None:
    """Append '<label> ' + a SEQ field (auto-number) to a caption paragraph."""
    paragraph.add_run(f"{label} ")

    p = paragraph._p
    begin_run = OxmlElement("w:r")
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    begin_run.append(fld_begin)
    p.append(begin_run)

    instr_run = OxmlElement("w:r")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = f" SEQ {label} \\* ARABIC "
    instr_run.append(instr)
    p.append(instr_run)

    sep_run = OxmlElement("w:r")
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    sep_run.append(fld_sep)
    p.append(sep_run)

    num_run = OxmlElement("w:r")
    num_t = OxmlElement("w:t")
    num_t.text = "1"
    num_run.append(num_t)
    p.append(num_run)

    end_run = OxmlElement("w:r")
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    end_run.append(fld_end)
    p.append(end_run)


async def add_caption(
    filename: str,
    text: str,
    label: str = "Figure",
    paragraph_index: Optional[int] = None,
    position: str = "after",
) -> str:
    """Add an auto-numbered caption (SEQ field) such as 'Figure 1: ...'.

    The caption uses a Word SEQ field so numbering stays correct as captions are
    added or removed, and can be collected by ``add_table_of_figures``.

    Args:
        filename: Path to the Word document.
        text: Caption description (appended after the number).
        label: Caption label / SEQ category ('Figure', 'Table', 'Equation', ...).
        paragraph_index: Anchor paragraph (e.g. the image/table paragraph).
            If omitted the caption is appended at the end of the document.
        position: 'after' (default) or 'before' the anchor paragraph.
    """
    filename = ensure_docx_extension(filename)
    if not os.path.exists(filename):
        return f"Document {filename} does not exist"

    is_writeable, error_message = check_file_writeable(filename)
    if not is_writeable:
        return f"Cannot modify document: {error_message}. Consider creating a copy first."

    if position not in {"after", "before"}:
        return "Invalid position. Use 'after' or 'before'."

    try:
        doc = Document(filename)

        if paragraph_index is not None and (
            paragraph_index < 0 or paragraph_index >= len(doc.paragraphs)
        ):
            return (
                f"Invalid paragraph index. Document has {len(doc.paragraphs)} "
                f"paragraphs (0-{len(doc.paragraphs) - 1})."
            )
        anchor = (
            doc.paragraphs[paragraph_index]._element
            if paragraph_index is not None
            else None
        )

        caption_para = doc.add_paragraph()
        try:
            caption_para.style = "Caption"
        except KeyError:
            pass

        _append_seq_caption_runs(caption_para, label)
        if text:
            caption_para.add_run(f": {text}")

        if anchor is not None:
            if position == "after":
                anchor.addnext(caption_para._element)
            else:
                anchor.addprevious(caption_para._element)

        doc.save(filename)
        return f"{label} caption added to {filename}."
    except Exception as e:
        return f"Failed to add caption: {str(e)}"


async def add_table_of_figures(
    filename: str,
    label: str = "Figure",
    title: Optional[str] = "Table of Figures",
    insert_at_start: bool = False,
    add_page_break_after: bool = False,
) -> str:
    """Insert a native Word Table of Figures field for a caption label.

    Collects all captions created with the matching ``label`` (via SEQ fields)
    into a refreshable list. Use after adding captions with ``add_caption``.
    """
    filename = ensure_docx_extension(filename)
    if not os.path.exists(filename):
        return f"Document {filename} does not exist"

    is_writeable, error_message = check_file_writeable(filename)
    if not is_writeable:
        return f"Cannot modify document: {error_message}. Consider creating a copy first."

    try:
        doc = Document(filename)
        _ensure_update_fields_on_open(doc)

        title_para = doc.add_paragraph(title) if title else None
        if title_para is not None and title:
            try:
                title_para.style = "TOC Heading"
            except KeyError:
                if title_para.runs:
                    title_para.runs[0].bold = True

        tof_para = doc.add_paragraph()
        instruction = f'TOC \\h \\z \\c "{label}"'
        _append_field_code_run(
            tof_para,
            instruction,
            display_text="Right-click to update field.",
        )

        if add_page_break_after:
            page_break_para = doc.add_paragraph()
            page_break_para.add_run().add_break(WD_BREAK.PAGE)
        else:
            page_break_para = None

        if insert_at_start and doc.paragraphs:
            first_element = doc.paragraphs[0]._element
            ordered = [p for p in [title_para, tof_para, page_break_para] if p is not None]
            for para in ordered:
                _insert_paragraph_before(para, first_element)

        doc.save(filename)
        return f"Table of figures (label '{label}') inserted into {filename}."
    except Exception as e:
        return f"Failed to insert table of figures: {str(e)}"
