"""
Comment authoring for WaDocx MCP.

python-docx 1.2.0 can *create* a comment (``Document.add_comment``) but has no
support for threaded replies or the resolved/"done" state. This module adds:

* locating the run(s) a comment should anchor to (by text or paragraph),
* creating a comment,
* replying to an existing comment (Word threads via ``commentsExtended.xml``),
* marking a comment thread resolved (the ``w15:done`` flag).
"""
from typing import List, Optional

from docx.oxml.ns import qn, nsmap
from docx.oxml import OxmlElement


W15 = "http://schemas.microsoft.com/office/word/2012/wordml"
W14 = "http://schemas.microsoft.com/office/word/2010/wordml"


def _w15(tag: str) -> str:
    return f"{{{W15}}}{tag}"


def _w14(tag: str) -> str:
    return f"{{{W14}}}{tag}"


def find_runs_for_text(doc, search_text: str, paragraph_index: Optional[int] = None,
                       occurrence: int = 1):
    """Return the list of Run objects that should anchor a comment.

    If ``paragraph_index`` is given and ``search_text`` is empty, anchors on the
    whole paragraph. Otherwise finds the paragraph containing the Nth
    ``occurrence`` of ``search_text`` and returns the runs spanning it.
    """
    paragraphs = list(doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                paragraphs.extend(cell.paragraphs)

    if paragraph_index is not None:
        if paragraph_index < 0 or paragraph_index >= len(doc.paragraphs):
            return None, "Invalid paragraph index."
        para = doc.paragraphs[paragraph_index]
        if not search_text:
            runs = para.runs
            if not runs:
                run = para.add_run("")
                runs = [run]
            return list(runs), None
        target_paras = [para]
    else:
        target_paras = paragraphs

    seen = 0
    for para in target_paras:
        if search_text and search_text in para.text:
            seen += 1
            if seen < occurrence:
                continue
            runs = para.runs
            if not runs:
                continue
            # Anchor on the runs whose combined text covers the match.
            offset = 0
            start_idx = end_idx = None
            match_start = para.text.find(search_text)
            match_end = match_start + len(search_text)
            for i, run in enumerate(runs):
                run_start = offset
                run_end = offset + len(run.text)
                if start_idx is None and run_end > match_start:
                    start_idx = i
                if run_start < match_end:
                    end_idx = i
                offset = run_end
            if start_idx is None:
                start_idx = 0
            if end_idx is None:
                end_idx = len(runs) - 1
            return list(runs[start_idx:end_idx + 1]), None

    return None, f"Text '{search_text}' not found."


def create_comment(doc, runs, text: str, author: str = "WaDocx",
                   initials: str = "WD"):
    """Create a comment anchored to ``runs`` and return it."""
    return doc.add_comment(runs, text=text, author=author, initials=initials or "")


def _comments_part(doc):
    """Return the comments part of the document, or None."""
    for rel in doc.part.rels.values():
        if rel.reltype.endswith("/comments"):
            return rel.target_part
    return None


def _ensure_para_id(paragraph_element) -> str:
    """Ensure a w:p has a w14:paraId, returning it."""
    existing = paragraph_element.get(_w14("paraId"))
    if existing:
        return existing
    # Derive a stable-ish 8-hex id from the element's position + id pool.
    import binascii

    raw = binascii.hexlify(repr(id(paragraph_element)).encode()).decode()
    para_id = raw[-8:].upper().rjust(8, "0")
    paragraph_element.set(_w14("paraId"), para_id)
    return para_id


def _get_or_create_comments_extended(doc):
    """Return the commentsExtended XML root, creating the part if needed."""
    from docx.opc.part import Part
    from docx.opc.packuri import PackURI

    main = doc.part
    reltype = "http://schemas.openxmlformats.org/officeDocument/2006/relationships/commentsExtended"
    content_type = (
        "application/vnd.openxmlformats-officedocument.wordprocessingml."
        "commentsExtended+xml"
    )
    for rel in main.rels.values():
        if rel.reltype == reltype:
            part = rel.target_part
            from lxml import etree

            return etree.fromstring(part.blob), part

    # Create a new commentsExtended part.
    from lxml import etree

    xml = (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>\n'
        '<w15:commentsEx xmlns:w15="%s" '
        'xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>'
        % W15
    )
    partname = PackURI("/word/commentsExtended.xml")
    part = Part(partname, content_type, xml.encode("utf-8"), main.package)
    main.relate_to(part, reltype)
    return etree.fromstring(part.blob), part


def reply_to_comment(doc, parent_comment_id: int, text: str,
                     author: str = "WaDocx", initials: str = "WD"):
    """Create a reply threaded under an existing comment.

    Anchors the reply on the same range as the parent comment and links them via
    ``commentsExtended.xml`` so Word displays a threaded conversation.
    """
    comments_part = _comments_part(doc)
    if comments_part is None:
        return None, "Document has no comments to reply to."

    parent_el = None
    for c in comments_part.element.findall(qn("w:comment")):
        if c.get(qn("w:id")) == str(parent_comment_id):
            parent_el = c
            break
    if parent_el is None:
        return None, f"Comment id {parent_comment_id} not found."

    # Find the runs the parent comment is anchored to.
    body = doc.element.body
    start_marker = None
    for el in body.iter(qn("w:commentRangeStart")):
        if el.get(qn("w:id")) == str(parent_comment_id):
            start_marker = el
            break
    runs = None
    if start_marker is not None:
        end_marker = None
        for el in body.iter(qn("w:commentRangeEnd")):
            if el.get(qn("w:id")) == str(parent_comment_id):
                end_marker = el
                break
        # Collect w:r elements between the markers.
        collected = []
        capture = False
        for el in body.iter():
            if el is start_marker:
                capture = True
                continue
            if el is end_marker:
                break
            if capture and el.tag == qn("w:r"):
                collected.append(el)
        if collected:
            from docx.text.run import Run

            runs = [Run(r, None) for r in collected]
    if not runs:
        return None, "Could not locate the parent comment's anchor range."

    reply = doc.add_comment(runs, text=text, author=author, initials=initials or "")

    # Link the reply under the parent in commentsExtended.
    parent_para_id = _ensure_para_id(parent_el.findall(qn("w:p"))[0])
    reply_el = None
    for c in comments_part.element.findall(qn("w:comment")):
        if c.get(qn("w:id")) == str(reply.comment_id):
            reply_el = c
            break
    if reply_el is not None:
        reply_para_id = _ensure_para_id(reply_el.findall(qn("w:p"))[0])
        root, part = _get_or_create_comments_extended(doc)
        ex = root.makeelement(_w15("commentEx"), {})
        ex.set(_w15("paraId"), reply_para_id)
        ex.set(_w15("paraIdParent"), parent_para_id)
        ex.set(_w15("done"), "0")
        root.append(ex)
        from lxml import etree

        part._blob = etree.tostring(root, xml_declaration=True, encoding="UTF-8",
                                    standalone=True)

    return reply, None


def resolve_comment(doc, comment_id: int, done: bool = True):
    """Mark a comment thread resolved/unresolved via the w15:done flag."""
    comments_part = _comments_part(doc)
    if comments_part is None:
        return False, "Document has no comments."
    target_el = None
    for c in comments_part.element.findall(qn("w:comment")):
        if c.get(qn("w:id")) == str(comment_id):
            target_el = c
            break
    if target_el is None:
        return False, f"Comment id {comment_id} not found."

    para_id = _ensure_para_id(target_el.findall(qn("w:p"))[0])
    root, part = _get_or_create_comments_extended(doc)
    found = None
    for ex in root.findall(_w15("commentEx")):
        if ex.get(_w15("paraId")) == para_id:
            found = ex
            break
    if found is None:
        found = root.makeelement(_w15("commentEx"), {})
        found.set(_w15("paraId"), para_id)
        root.append(found)
    found.set(_w15("done"), "1" if done else "0")
    from lxml import etree

    part._blob = etree.tostring(root, xml_declaration=True, encoding="UTF-8",
                                standalone=True)
    return True, None
