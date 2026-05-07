"""
Document utility functions for WaDocx MCP.
"""
import json
import os
from typing import Dict, List, Any, Optional
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Inches
from docx.text.paragraph import Paragraph


def apply_block_alignment(paragraph, alignment: Optional[str]) -> None:
    """Apply a simple alignment keyword to a paragraph."""
    if not alignment:
        return

    alignment_map = {
        "left": WD_ALIGN_PARAGRAPH.LEFT,
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "right": WD_ALIGN_PARAGRAPH.RIGHT,
        "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
    }
    paragraph.alignment = alignment_map.get(alignment.lower())


def get_document_properties(doc_path: str) -> Dict[str, Any]:
    """Get properties of a Word document."""
    import os
    if not os.path.exists(doc_path):
        return {"error": f"Document {doc_path} does not exist"}
    
    try:
        doc = Document(doc_path)
        core_props = doc.core_properties
        
        return {
            "title": core_props.title or "",
            "author": core_props.author or "",
            "subject": core_props.subject or "",
            "keywords": core_props.keywords or "",
            "created": str(core_props.created) if core_props.created else "",
            "modified": str(core_props.modified) if core_props.modified else "",
            "last_modified_by": core_props.last_modified_by or "",
            "revision": core_props.revision or 0,
            "page_count": len(doc.sections),
            "word_count": sum(len(paragraph.text.split()) for paragraph in doc.paragraphs),
            "paragraph_count": len(doc.paragraphs),
            "table_count": len(doc.tables)
        }
    except Exception as e:
        return {"error": f"Failed to get document properties: {str(e)}"}


def extract_document_text(doc_path: str) -> str:
    """Extract all text from a Word document."""
    import os
    if not os.path.exists(doc_path):
        return f"Document {doc_path} does not exist"
    
    try:
        doc = Document(doc_path)
        text = []
        
        for paragraph in doc.paragraphs:
            text.append(paragraph.text)
            
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        text.append(paragraph.text)
        
        return "\n".join(text)
    except Exception as e:
        return f"Failed to extract text: {str(e)}"


def get_document_structure(doc_path: str) -> Dict[str, Any]:
    """Get the structure of a Word document."""
    import os
    if not os.path.exists(doc_path):
        return {"error": f"Document {doc_path} does not exist"}
    
    try:
        doc = Document(doc_path)
        structure = {
            "paragraphs": [],
            "tables": []
        }
        
        # Get paragraphs
        for i, para in enumerate(doc.paragraphs):
            structure["paragraphs"].append({
                "index": i,
                "text": para.text[:100] + ("..." if len(para.text) > 100 else ""),
                "style": para.style.name if para.style else "Normal"
            })
        
        # Get tables
        for i, table in enumerate(doc.tables):
            table_data = {
                "index": i,
                "rows": len(table.rows),
                "columns": len(table.columns),
                "preview": []
            }
            
            # Get sample of table data
            max_rows = min(3, len(table.rows))
            for row_idx in range(max_rows):
                row_data = []
                max_cols = min(3, len(table.columns))
                for col_idx in range(max_cols):
                    try:
                        cell_text = table.cell(row_idx, col_idx).text
                        row_data.append(cell_text[:20] + ("..." if len(cell_text) > 20 else ""))
                    except IndexError:
                        row_data.append("N/A")
                table_data["preview"].append(row_data)
            
            structure["tables"].append(table_data)
        
        return structure
    except Exception as e:
        return {"error": f"Failed to get document structure: {str(e)}"}


def find_paragraph_by_text(doc, text, partial_match=False):
    """
    Find paragraphs containing specific text.
    
    Args:
        doc: Document object
        text: Text to search for
        partial_match: If True, matches paragraphs containing the text; if False, matches exact text
        
    Returns:
        List of paragraph indices that match the criteria
    """
    matching_paragraphs = []
    
    for i, para in enumerate(doc.paragraphs):
        if partial_match and text in para.text:
            matching_paragraphs.append(i)
        elif not partial_match and para.text == text:
            matching_paragraphs.append(i)
            
    return matching_paragraphs


def find_and_replace_text(doc, old_text, new_text):
    """
    Find and replace text throughout the document, skipping Table of Contents (TOC) paragraphs.
    
    Args:
        doc: Document object
        old_text: Text to find
        new_text: Text to replace with
        
    Returns:
        Number of replacements made
    """
    count = 0
    
    # Search in paragraphs
    for para in doc.paragraphs:
        # Skip TOC paragraphs
        if para.style and para.style.name.startswith("TOC"):
            continue
        if old_text in para.text:
            for run in para.runs:
                if old_text in run.text:
                    run.text = run.text.replace(old_text, new_text)
                    count += 1
    
    # Search in tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    # Skip TOC paragraphs in tables
                    if para.style and para.style.name.startswith("TOC"):
                        continue
                    if old_text in para.text:
                        for run in para.runs:
                            if old_text in run.text:
                                run.text = run.text.replace(old_text, new_text)
                                count += 1
    
    return count


def get_document_xml(doc_path: str) -> str:
    """Extract and return the raw XML structure of the Word document (word/document.xml)."""
    import os
    import zipfile
    if not os.path.exists(doc_path):
        return f"Document {doc_path} does not exist"
    try:
        with zipfile.ZipFile(doc_path) as docx_zip:
            with docx_zip.open('word/document.xml') as xml_file:
                return xml_file.read().decode('utf-8')
    except Exception as e:
        return f"Failed to extract XML: {str(e)}"


def insert_header_near_text(doc_path: str, target_text: str = None, header_title: str = "", position: str = 'after', header_style: str = 'Heading 1', target_paragraph_index: int = None) -> str:
    """Insert a header (with specified style) before or after the target paragraph. Specify by text or paragraph index. Skips TOC paragraphs in text search."""
    import os
    from docx import Document
    if not os.path.exists(doc_path):
        return f"Document {doc_path} does not exist"
    try:
        doc = Document(doc_path)
        found = False
        para = None
        if target_paragraph_index is not None:
            if target_paragraph_index < 0 or target_paragraph_index >= len(doc.paragraphs):
                return f"Invalid target_paragraph_index: {target_paragraph_index}. Document has {len(doc.paragraphs)} paragraphs."
            para = doc.paragraphs[target_paragraph_index]
            found = True
        else:
            for i, p in enumerate(doc.paragraphs):
                # Skip TOC paragraphs
                if p.style and p.style.name.lower().startswith("toc"):
                    continue
                if target_text and target_text in p.text:
                    para = p
                    found = True
                    break
        if not found or para is None:
            return f"Target paragraph not found (by index or text). (TOC paragraphs are skipped in text search)"
        # Save anchor index before insertion
        if target_paragraph_index is not None:
            anchor_index = target_paragraph_index
        else:
            anchor_index = None
            for i, p in enumerate(doc.paragraphs):
                if p is para:
                    anchor_index = i
                    break
        new_para = doc.add_paragraph(header_title, style=header_style)
        if position == 'before':
            para._element.addprevious(new_para._element)
        else:
            para._element.addnext(new_para._element)
        doc.save(doc_path)
        if anchor_index is not None:
            return f"Header '{header_title}' (style: {header_style}) inserted {position} paragraph (index {anchor_index})."
        else:
            return f"Header '{header_title}' (style: {header_style}) inserted {position} the target paragraph."
    except Exception as e:
        return f"Failed to insert header: {str(e)}"


def insert_line_or_paragraph_near_text(doc_path: str, target_text: str = None, line_text: str = "", position: str = 'after', line_style: str = None, target_paragraph_index: int = None) -> str:
    """
    Insert a new line or paragraph (with specified or matched style) before or after the target paragraph.
    You can specify the target by text (first match) or by paragraph index.
    Skips paragraphs whose style name starts with 'TOC' if using text search.
    """
    import os
    from docx import Document
    if not os.path.exists(doc_path):
        return f"Document {doc_path} does not exist"
    try:
        doc = Document(doc_path)
        found = False
        para = None
        if target_paragraph_index is not None:
            if target_paragraph_index < 0 or target_paragraph_index >= len(doc.paragraphs):
                return f"Invalid target_paragraph_index: {target_paragraph_index}. Document has {len(doc.paragraphs)} paragraphs."
            para = doc.paragraphs[target_paragraph_index]
            found = True
        else:
            for i, p in enumerate(doc.paragraphs):
                # Skip TOC paragraphs
                if p.style and p.style.name.lower().startswith("toc"):
                    continue
                if target_text and target_text in p.text:
                    para = p
                    found = True
                    break
        if not found or para is None:
            return f"Target paragraph not found (by index or text). (TOC paragraphs are skipped in text search)"
        # Save anchor index before insertion
        if target_paragraph_index is not None:
            anchor_index = target_paragraph_index
        else:
            anchor_index = None
            for i, p in enumerate(doc.paragraphs):
                if p is para:
                    anchor_index = i
                    break
        # Determine style: use provided or match target
        style = line_style if line_style else para.style
        new_para = doc.add_paragraph(line_text, style=style)
        if position == 'before':
            para._element.addprevious(new_para._element)
        else:
            para._element.addnext(new_para._element)
        doc.save(doc_path)
        if anchor_index is not None:
            return f"Line/paragraph inserted {position} paragraph (index {anchor_index}) with style '{style}'."
        else:
            return f"Line/paragraph inserted {position} the target paragraph with style '{style}'."
    except Exception as e:
        return f"Failed to insert line/paragraph: {str(e)}"


def add_bullet_numbering(paragraph, num_id=1, level=0):
    """
    Add bullet/numbering XML to a paragraph.

    Args:
        paragraph: python-docx Paragraph object
        num_id: Numbering definition ID (1=bullets, 2=numbers, etc.)
        level: Indentation level (0=first level, 1=second level, etc.)

    Returns:
        The modified paragraph
    """
    # Get or create paragraph properties
    pPr = paragraph._element.get_or_add_pPr()

    # Remove existing numPr if any (to avoid duplicates)
    existing_numPr = pPr.find(qn('w:numPr'))
    if existing_numPr is not None:
        pPr.remove(existing_numPr)

    # Create numbering properties element
    numPr = OxmlElement('w:numPr')

    # Set indentation level
    ilvl = OxmlElement('w:ilvl')
    ilvl.set(qn('w:val'), str(level))
    numPr.append(ilvl)

    # Set numbering definition ID
    numId = OxmlElement('w:numId')
    numId.set(qn('w:val'), str(num_id))
    numPr.append(numId)

    # Add to paragraph properties
    pPr.append(numPr)

    return paragraph


def insert_numbered_list_near_text(doc_path: str, target_text: str = None, list_items: list = None, position: str = 'after', target_paragraph_index: int = None, bullet_type: str = 'bullet') -> str:
    """
    Insert a bulleted or numbered list before or after the target paragraph. Specify by text or paragraph index. Skips TOC paragraphs in text search.
    Args:
        doc_path: Path to the Word document
        target_text: Text to search for in paragraphs (optional if using index)
        list_items: List of strings, each as a list item
        position: 'before' or 'after' (default: 'after')
        target_paragraph_index: Optional paragraph index to use as anchor
        bullet_type: 'bullet' for bullets (•), 'number' for numbers (1,2,3) (default: 'bullet')
    Returns:
        Status message
    """
    import os
    from docx import Document
    if not os.path.exists(doc_path):
        return f"Document {doc_path} does not exist"
    try:
        doc = Document(doc_path)
        found = False
        para = None
        if target_paragraph_index is not None:
            if target_paragraph_index < 0 or target_paragraph_index >= len(doc.paragraphs):
                return f"Invalid target_paragraph_index: {target_paragraph_index}. Document has {len(doc.paragraphs)} paragraphs."
            para = doc.paragraphs[target_paragraph_index]
            found = True
        else:
            for i, p in enumerate(doc.paragraphs):
                # Skip TOC paragraphs
                if p.style and p.style.name.lower().startswith("toc"):
                    continue
                if target_text and target_text in p.text:
                    para = p
                    found = True
                    break
        if not found or para is None:
            return f"Target paragraph not found (by index or text). (TOC paragraphs are skipped in text search)"
        # Save anchor index before insertion
        if target_paragraph_index is not None:
            anchor_index = target_paragraph_index
        else:
            anchor_index = None
            for i, p in enumerate(doc.paragraphs):
                if p is para:
                    anchor_index = i
                    break
        # Determine numbering ID based on bullet_type
        num_id = 1 if bullet_type == 'bullet' else 2

        # Use ListParagraph style for proper list formatting
        style_name = None
        for candidate in ['List Paragraph', 'ListParagraph', 'Normal']:
            try:
                _ = doc.styles[candidate]
                style_name = candidate
                break
            except KeyError:
                continue
        if not style_name:
            style_name = None  # fallback to default

        new_paras = []
        for item in (list_items or []):
            p = doc.add_paragraph(item, style=style_name)
            # Add bullet numbering XML - this is the fix!
            add_bullet_numbering(p, num_id=num_id, level=0)
            new_paras.append(p)
        # Move the new paragraphs to the correct position
        for p in reversed(new_paras):
            if position == 'before':
                para._element.addprevious(p._element)
            else:
                para._element.addnext(p._element)
        doc.save(doc_path)
        list_type = "bulleted" if bullet_type == 'bullet' else "numbered"
        if anchor_index is not None:
            return f"{list_type.capitalize()} list with {len(new_paras)} items inserted {position} paragraph (index {anchor_index})."
        else:
            return f"{list_type.capitalize()} list with {len(new_paras)} items inserted {position} the target paragraph."
    except Exception as e:
        return f"Failed to insert numbered list: {str(e)}"


def is_toc_paragraph(para):
    """Devuelve True si el párrafo tiene un estilo de tabla de contenido (TOC)."""
    return para.style and para.style.name.upper().startswith("TOC")


def is_heading_paragraph(para):
    """Devuelve True si el párrafo tiene un estilo de encabezado (Heading 1, Heading 2, etc)."""
    return para.style and para.style.name.lower().startswith("heading")


# --- Helper: Get style name from a <w:p> element ---
def get_paragraph_style(el):
    from docx.oxml.ns import qn
    pPr = el.find(qn('w:pPr'))
    if pPr is not None:
        pStyle = pPr.find(qn('w:pStyle'))
        if pStyle is not None and 'w:val' in pStyle.attrib:
            return pStyle.attrib['w:val']
    return None


def normalize_paragraph_text(text: Optional[str]) -> str:
    """Normalize paragraph text for stable matching."""
    return " ".join((text or "").split()).strip()


def is_paragraph_element(el) -> bool:
    """Return True when the XML element is a paragraph."""
    return str(getattr(el, "tag", "")).endswith("}p")


def is_table_element(el) -> bool:
    """Return True when the XML element is a table."""
    return str(getattr(el, "tag", "")).endswith("}tbl")


def get_body_elements(doc) -> List[Any]:
    """Return body elements excluding section properties."""
    elements = []
    for el in doc.element.body.iterchildren():
        if el.tag == qn('w:sectPr'):
            continue
        elements.append(el)
    return elements


def get_element_text(el) -> str:
    """Extract visible text from a paragraph element."""
    if not is_paragraph_element(el):
        return ""
    text_tag = qn("w:t")
    return normalize_paragraph_text(
        "".join(node.text or "" for node in el.iter() if getattr(node, "tag", None) == text_tag)
    )


def get_paragraph_from_element(doc, el):
    """Wrap a paragraph XML node as a python-docx paragraph."""
    if not is_paragraph_element(el):
        return None
    for para in doc.paragraphs:
        if para._element == el:
            return para
    return None


def is_heading_or_toc_element(doc, el) -> bool:
    """Return True if the body element is a heading or TOC paragraph."""
    para = get_paragraph_from_element(doc, el)
    if para is None or not para.style:
        return False
    style_name = para.style.name.lower()
    return style_name.startswith(("heading", "título", "toc"))


def remove_body_elements(doc, start_idx: int, end_idx: int) -> int:
    """Remove a contiguous body-element range and return the count removed."""
    elements = get_body_elements(doc)
    removed = 0
    for el in elements[start_idx:end_idx]:
        parent = el.getparent()
        if parent is not None:
            parent.remove(el)
            removed += 1
    return removed


def insert_text_paragraph_after_element(doc, anchor_element, text: str, style_name: Optional[str] = None):
    """Insert a paragraph after a body element and return the new paragraph element."""
    new_para = doc.add_paragraph(text)
    if style_name:
        try:
            new_para.style = style_name
        except KeyError:
            pass
    anchor_element.addnext(new_para._element)
    return new_para._element


def ensure_update_fields_on_open(doc) -> None:
    """Ask Word to update fields when the document opens."""
    settings_element = doc.settings.element
    update_fields = settings_element.find(qn("w:updateFields"))
    if update_fields is None:
        update_fields = OxmlElement("w:updateFields")
        settings_element.append(update_fields)
    update_fields.set(qn("w:val"), "true")


def append_field_code_run(paragraph, instruction: str, display_text: str = "") -> None:
    """Append a native Word field code to a paragraph."""
    p = paragraph._p

    begin_run = OxmlElement("w:r")
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    fld_char_begin.set(qn("w:dirty"), "true")
    begin_run.append(fld_char_begin)
    p.append(begin_run)

    instr_run = OxmlElement("w:r")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = f" {instruction} "
    instr_run.append(instr_text)
    p.append(instr_run)

    separate_run = OxmlElement("w:r")
    fld_char_sep = OxmlElement("w:fldChar")
    fld_char_sep.set(qn("w:fldCharType"), "separate")
    separate_run.append(fld_char_sep)
    p.append(separate_run)

    if display_text:
        paragraph.add_run(display_text)

    end_run = OxmlElement("w:r")
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    end_run.append(fld_char_end)
    p.append(end_run)


def build_toc_instruction(max_level: int = 3, toc_style: str = "dotted") -> str:
    """Build a Word TOC field instruction for a supported visual style."""
    max_level = max(1, min(int(max_level), 9))
    normalized = (toc_style or "dotted").strip().lower().replace("-", "_")
    aliases = {
        "default": "dotted",
        "dots": "dotted",
        "dotted_leader": "dotted",
        "dotted_leaders": "dotted",
        "page_number": "page_numbers",
        "page_numbers": "page_numbers",
        "plain": "page_numbers",
        "plain_page_numbers": "page_numbers",
        "blue_links": "links",
        "blue_link": "links",
        "link": "links",
        "links": "links",
        "web": "links",
        "web_links": "links",
    }
    style = aliases.get(normalized, normalized)
    base = f'TOC \\o "1-{max_level}" \\h \\z \\u'

    if style == "page_numbers":
        return f'{base} \\p " "'
    if style == "links":
        return f'{base} \\n "1-{max_level}"'
    return base


def insert_live_toc_after_element(
    doc,
    anchor_element,
    title: Optional[str] = "Contents",
    max_level: int = 3,
    add_page_break_after: bool = False,
    toc_style: str = "dotted",
):
    """Insert a native Word TOC field after a body element."""
    ensure_update_fields_on_open(doc)
    max_level = max(1, min(int(max_level), 9))
    current_element = anchor_element
    inserted = 0

    if title:
        title_para = doc.add_paragraph(title)
        try:
            title_para.style = "TOC Heading"
        except KeyError:
            if title_para.runs:
                title_para.runs[0].bold = True
        current_element.addnext(title_para._element)
        current_element = title_para._element
        inserted += 1

    toc_para = doc.add_paragraph()
    append_field_code_run(
        toc_para,
        build_toc_instruction(max_level=max_level, toc_style=toc_style),
        display_text="Right-click to update field.",
    )
    current_element.addnext(toc_para._element)
    current_element = toc_para._element
    inserted += 1

    if add_page_break_after:
        page_break_para = doc.add_paragraph()
        page_break_para.add_run().add_break(WD_BREAK.PAGE)
        current_element.addnext(page_break_para._element)
        current_element = page_break_para._element
        inserted += 1

    return current_element, inserted


def insert_content_blocks_after_element(
    doc,
    anchor_element,
    blocks: List[Dict[str, Any]],
    default_paragraph_style: Optional[str] = "Normal"
) -> int:
    """
    Insert parsed content blocks after an anchor element.

    Supported block types: paragraph, heading, list, table, image, page_break, toc.
    Returns the number of body elements inserted.
    """
    current_element = anchor_element
    inserted = 0

    for block in blocks:
        block_type = block.get("type", "paragraph")

        if block_type == "heading":
            level = max(1, min(int(block.get("level", 1)), 9))
            new_para = doc.add_heading(block.get("text", ""), level=level)
            apply_block_alignment(new_para, block.get("alignment"))
            current_element.addnext(new_para._element)
            current_element = new_para._element
            inserted += 1
            continue

        if block_type == "toc":
            current_element, added = insert_live_toc_after_element(
                doc,
                current_element,
                title=block.get("title", "Contents"),
                max_level=block.get("max_level", 3),
                add_page_break_after=bool(block.get("add_page_break_after")),
                toc_style=block.get("toc_style", block.get("style", "dotted")),
            )
            inserted += added
            continue

        if block_type == "page_break":
            new_para = doc.add_paragraph("")
            new_para.add_run().add_break(WD_BREAK.PAGE)
            current_element.addnext(new_para._element)
            current_element = new_para._element
            inserted += 1
            continue

        if block_type == "list":
            ordered = bool(block.get("ordered"))
            style_name = "List Number" if ordered else "List Bullet"
            num_id = 2 if ordered else 1
            for item in block.get("items", []):
                new_para = doc.add_paragraph(item)
                try:
                    new_para.style = style_name
                except KeyError:
                    pass
                add_bullet_numbering(new_para, num_id=num_id, level=0)
                apply_block_alignment(new_para, block.get("alignment"))
                current_element.addnext(new_para._element)
                current_element = new_para._element
                inserted += 1
            continue

        if block_type == "table":
            rows = block.get("rows", [])
            if not rows:
                continue
            column_count = max(len(row) for row in rows)
            table = doc.add_table(rows=len(rows), cols=column_count)
            for row_idx, row in enumerate(rows):
                for col_idx, value in enumerate(row):
                    table.cell(row_idx, col_idx).text = value
            if block.get("alignment") == "center":
                table.alignment = WD_TABLE_ALIGNMENT.CENTER
            current_element.addnext(table._element)
            current_element = table._element
            inserted += 1
            continue

        if block_type == "image":
            image_path = block.get("path", "")
            if not image_path:
                continue
            if not os.path.exists(image_path):
                raise FileNotFoundError(f"Markdown image file not found: {image_path}")

            new_para = doc.add_paragraph()
            image_run = new_para.add_run()
            width = block.get("width")
            if width is not None:
                inline_shape = image_run.add_picture(image_path, width=Inches(width))
            else:
                inline_shape = image_run.add_picture(image_path)

            alt_text = block.get("alt", "").strip()
            if alt_text:
                doc_pr = inline_shape._inline.docPr
                doc_pr.set("name", alt_text[:255])
                doc_pr.set("descr", alt_text[:255])
                doc_pr.set("title", alt_text[:255])
            apply_block_alignment(new_para, block.get("alignment"))

            current_element.addnext(new_para._element)
            current_element = new_para._element
            inserted += 1
            continue

        current_element = insert_text_paragraph_after_element(
            doc,
            current_element,
            block.get("text", ""),
            block.get("style", default_paragraph_style)
        )
        apply_block_alignment(get_paragraph_from_element(doc, current_element), block.get("alignment"))
        inserted += 1

    return inserted

# --- Main: Delete everything under a header until next heading/TOC ---
def delete_block_under_header(doc, header_text):
    """
    Remove all elements (paragraphs, tables, etc.) after the header (by text) and before the next heading/TOC (by style).
    Returns: (header_element, elements_removed)
    """
    # Find the header paragraph by text (like delete_paragraph finds by index)
    header_para = None
    header_idx = None
    
    for i, para in enumerate(doc.paragraphs):
        if para.text.strip().lower() == header_text.strip().lower():
            header_para = para
            header_idx = i
            break
    
    if header_para is None:
        return None, 0
    
    # Find the next heading/TOC paragraph to determine the end of the block
    end_idx = None
    for i in range(header_idx + 1, len(doc.paragraphs)):
        para = doc.paragraphs[i]
        if para.style and para.style.name.lower().startswith(('heading', 'título', 'toc')):
            end_idx = i
            break
    
    # If no next heading found, delete until end of document
    if end_idx is None:
        end_idx = len(doc.paragraphs)
    
    # Remove paragraphs by index (like delete_paragraph does)
    removed_count = 0
    for i in range(header_idx + 1, end_idx):
        if i < len(doc.paragraphs):  # Safety check
            para = doc.paragraphs[header_idx + 1]  # Always remove the first paragraph after header
            p = para._p
            p.getparent().remove(p)
            removed_count += 1
    
    return header_para._p, removed_count

# --- Usage in replace_paragraph_block_below_header ---
def replace_paragraph_block_below_header(
    doc_path: str,
    header_text: str,
    new_paragraphs: list,
    detect_block_end_fn=None,
    new_paragraph_style: str = None
) -> str:
    """
    Reemplaza todo el contenido debajo de una cabecera (por texto), hasta el siguiente encabezado/TOC (por estilo).
    """
    from docx import Document
    import os
    if not os.path.exists(doc_path):
        return f"Document {doc_path} not found."
    
    doc = Document(doc_path)
    
    # Find the header paragraph first
    header_para = None
    header_idx = None
    for i, para in enumerate(doc.paragraphs):
        para_text = para.text.strip().lower()
        is_toc = is_toc_paragraph(para)
        if para_text == header_text.strip().lower() and not is_toc:
            header_para = para
            header_idx = i
            break
    
    if header_para is None:
        return f"Header '{header_text}' not found in document."
    
    # Delete everything under the header using the same document instance
    header_el, removed_count = delete_block_under_header(doc, header_text)
    
    # Now insert new paragraphs after the header (which should still be in the document)
    style_to_use = new_paragraph_style or "Normal"
    
    # Find the header again after deletion (it should still be there)
    current_para = header_para
    for text in new_paragraphs:
        new_para = doc.add_paragraph(text, style=style_to_use)
        current_para._element.addnext(new_para._element)
        current_para = new_para
    
    doc.save(doc_path)
    return f"Replaced content under '{header_text}' with {len(new_paragraphs)} paragraph(s), style: {style_to_use}, removed {removed_count} elements."


def replace_block_between_manual_anchors(
    doc_path: str,
    start_anchor_text: str,
    new_paragraphs: list,
    end_anchor_text: str = None,
    match_fn=None,
    new_paragraph_style: str = None
) -> str:
    """
    Replace all content (paragraphs, tables, etc.) between start_anchor_text and end_anchor_text (or next logical header if not provided).
    If end_anchor_text is None, deletes until next visually distinct paragraph (bold, all caps, or different font size), or end of document.
    Inserts new_paragraphs after the start anchor.
    """
    from docx import Document
    import os
    if not os.path.exists(doc_path):
        return f"Document {doc_path} not found."
    doc = Document(doc_path)
    body = doc.element.body
    elements = list(body)
    start_idx = None
    end_idx = None
    # Find start anchor
    for i, el in enumerate(elements):
        if el.tag == CT_P.tag:
            p_text = "".join([node.text or '' for node in el.iter() if node.tag.endswith('}t')]).strip()
            if match_fn:
                if match_fn(p_text, el):
                    start_idx = i
                    break
            elif p_text == start_anchor_text.strip():
                start_idx = i
                break
    if start_idx is None:
        return f"Start anchor '{start_anchor_text}' not found."
    # Find end anchor
    if end_anchor_text:
        for i in range(start_idx + 1, len(elements)):
            el = elements[i]
            if el.tag == CT_P.tag:
                p_text = "".join([node.text or '' for node in el.iter() if node.tag.endswith('}t')]).strip()
                if match_fn:
                    if match_fn(p_text, el, is_end=True):
                        end_idx = i
                        break
                elif p_text == end_anchor_text.strip():
                    end_idx = i
                    break
    else:
        # Heuristic: next visually distinct paragraph (bold, all caps, or different font size), or end of document
        for i in range(start_idx + 1, len(elements)):
            el = elements[i]
            if el.tag == CT_P.tag:
                # Check for bold, all caps, or font size
                runs = [node for node in el.iter() if node.tag.endswith('}r')]
                for run in runs:
                    rpr = run.find(qn('w:rPr'))
                    if rpr is not None:
                        if rpr.find(qn('w:b')) is not None or rpr.find(qn('w:caps')) is not None or rpr.find(qn('w:sz')) is not None:
                            end_idx = i
                            break
                if end_idx is not None:
                    break
    # Mark elements for removal
    to_remove = []
    for i in range(start_idx + 1, end_idx if end_idx is not None else len(elements)):
        to_remove.append(elements[i])
    for el in to_remove:
        body.remove(el)
    doc.save(doc_path)
    # Reload and find start anchor for insertion
    doc = Document(doc_path)
    paras = doc.paragraphs
    anchor_idx = None
    for i, para in enumerate(paras):
        if para.text.strip() == start_anchor_text.strip():
            anchor_idx = i
            break
    if anchor_idx is None:
        return f"Start anchor '{start_anchor_text}' not found after deletion (unexpected)."
    anchor_para = paras[anchor_idx]
    style_to_use = new_paragraph_style or "Normal"
    for text in new_paragraphs:
        new_para = doc.add_paragraph(text, style=style_to_use)
        anchor_para._element.addnext(new_para._element)
        anchor_para = new_para
    doc.save(doc_path)
    return f"Replaced content between '{start_anchor_text}' and '{end_anchor_text or 'next logical header'}' with {len(new_paragraphs)} paragraph(s), style: {style_to_use}, removed {len(to_remove)} elements."


# Stable body-element-based replacements. These later definitions intentionally
# override the earlier paragraph-index implementations above.
def delete_block_under_header(doc, header_text):
    """
    Remove all body elements after a header paragraph and before the next
    heading/TOC block. Returns (header_element, elements_removed).
    """
    header_text_normalized = normalize_paragraph_text(header_text).lower()
    elements = get_body_elements(doc)
    header_idx = None
    header_element = None

    for i, el in enumerate(elements):
        if not is_paragraph_element(el):
            continue
        para = get_paragraph_from_element(doc, el)
        if para is None or is_toc_paragraph(para):
            continue
        if normalize_paragraph_text(para.text).lower() == header_text_normalized:
            header_idx = i
            header_element = el
            break

    if header_idx is None:
        return None, 0

    end_idx = len(elements)
    for i in range(header_idx + 1, len(elements)):
        if is_heading_or_toc_element(doc, elements[i]):
            end_idx = i
            break

    removed_count = remove_body_elements(doc, header_idx + 1, end_idx)
    return header_element, removed_count


def replace_paragraph_block_below_header(
    doc_path: str,
    header_text: str,
    new_paragraphs: list,
    detect_block_end_fn=None,
    new_paragraph_style: str = None
) -> str:
    """
    Replace all body content beneath a header until the next heading or TOC.
    """
    from docx import Document
    import os
    if not os.path.exists(doc_path):
        return f"Document {doc_path} not found."

    doc = Document(doc_path)
    header_el, removed_count = delete_block_under_header(doc, header_text)
    if header_el is None:
        return f"Header '{header_text}' not found in document."

    style_to_use = new_paragraph_style or "Normal"
    current_element = header_el
    inserted = 0
    for text in new_paragraphs:
        current_element = insert_text_paragraph_after_element(doc, current_element, text, style_to_use)
        inserted += 1

    doc.save(doc_path)
    return (
        f"Replaced content under '{header_text}' with {inserted} paragraph(s), "
        f"style: {style_to_use}, removed {removed_count} body element(s)."
    )


def replace_block_between_manual_anchors(
    doc_path: str,
    start_anchor_text: str,
    new_paragraphs: list,
    end_anchor_text: str = None,
    match_fn=None,
    new_paragraph_style: str = None
) -> str:
    """
    Replace all body content between anchor paragraphs.

    If end_anchor_text is omitted, replacement stops at the next heading/TOC
    or the end of the document body.
    """
    from docx import Document
    import os
    if not os.path.exists(doc_path):
        return f"Document {doc_path} not found."

    doc = Document(doc_path)
    elements = get_body_elements(doc)
    start_idx = None
    end_idx = None
    start_anchor_normalized = normalize_paragraph_text(start_anchor_text)
    end_anchor_normalized = normalize_paragraph_text(end_anchor_text) if end_anchor_text else None

    for i, el in enumerate(elements):
        if not is_paragraph_element(el):
            continue
        p_text = get_element_text(el)
        if match_fn:
            if match_fn(p_text, el):
                start_idx = i
                break
        elif p_text == start_anchor_normalized:
            start_idx = i
            break

    if start_idx is None:
        return f"Start anchor '{start_anchor_text}' not found."

    if end_anchor_text:
        for i in range(start_idx + 1, len(elements)):
            el = elements[i]
            if not is_paragraph_element(el):
                continue
            p_text = get_element_text(el)
            if match_fn:
                if match_fn(p_text, el, is_end=True):
                    end_idx = i
                    break
            elif p_text == end_anchor_normalized:
                end_idx = i
                break
    else:
        for i in range(start_idx + 1, len(elements)):
            if is_heading_or_toc_element(doc, elements[i]):
                end_idx = i
                break

    removed_count = remove_body_elements(
        doc,
        start_idx + 1,
        end_idx if end_idx is not None else len(elements)
    )

    current_element = get_body_elements(doc)[start_idx]
    style_to_use = new_paragraph_style or "Normal"
    inserted = 0
    for text in new_paragraphs:
        current_element = insert_text_paragraph_after_element(doc, current_element, text, style_to_use)
        inserted += 1

    doc.save(doc_path)
    return (
        f"Replaced content between '{start_anchor_text}' and "
        f"'{end_anchor_text or 'next heading/TOC/end'}' with {inserted} paragraph(s), "
        f"style: {style_to_use}, removed {removed_count} body element(s)."
    )

