"""
Markdown import/export utilities for WaDocx MCP.
"""
import base64
import hashlib
import os
import re
import shutil
import textwrap
from copy import deepcopy
from typing import Any, Dict, List, Optional

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml.ns import qn
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import Table

from word_document_server.utils.document_utils import (
    delete_block_under_header,
    get_body_elements,
    get_paragraph_from_element,
    is_paragraph_element,
    is_table_element,
    insert_content_blocks_after_element,
    normalize_paragraph_text,
)


HEADING_RE = re.compile(r"^(#{1,6})\s+(.*)$")
ORDERED_LIST_RE = re.compile(r"^\s*(\d+)\.\s+(.*)$")
BULLET_LIST_RE = re.compile(r"^\s*[-*+]\s+(.*)$")
IMAGE_RE = re.compile(r'^!\[(?P<alt>.*?)\]\((?P<path><.*?>|[^)]+)\)$')
IMAGE_LINK_RE = re.compile(r'!\[(?P<alt>(?:\\\]|[^\]])*)\]\(<(?P<path>[^>]+)>\)')
PAGE_BREAK_RE = re.compile(r"^<!--\s*PAGE BREAK\s*-->$")
SECTION_BREAK_RE = re.compile(r"^<!--\s*SECTION BREAK\s*-->$")
TOC_RE = re.compile(r"^<!--\s*(?:TOC|wadocx:toc)(?P<body>.*?)-->\s*$", re.IGNORECASE | re.DOTALL)
DIV_OPEN_RE = re.compile(r'^<div\s+align="(?P<align>left|center|right|justify)"\s*>$', re.IGNORECASE)
DIV_CLOSE_RE = re.compile(r"^</div>$", re.IGNORECASE)
FIDELITY_RE = re.compile(
    r"<!--\s*wadocx:fidelity-bundle\s*\n(?P<body>.*?)\n-->\s*",
    re.DOTALL,
)
BASE_TEMPLATE_RE = re.compile(
    r"<!--\s*wadocx:base-template-md\s*\n(?P<body>.*?)\n-->\s*",
    re.DOTALL,
)

FIDELITY_VERSION = "1"
REPLACE_MODES = {
    "auto",
    "exact_restore",
    "editable_rebuild",
    "editable_rebuild_with_template",
}


def _normalize_inline_markdown_text(text: str) -> str:
    """Collapse simple inline markdown formatting to plain text."""
    text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
    text = re.sub(r"(?<!\*)\*(?!\*)(.+?)(?<!\*)\*(?!\*)", r"\1", text)
    text = re.sub(r"`([^`]+)`", r"\1", text)
    return text.strip()


def _is_markdown_table_line(line: str) -> bool:
    stripped = line.strip()
    return stripped.startswith("|") and stripped.endswith("|") and stripped.count("|") >= 2


def _split_markdown_row(line: str) -> List[str]:
    return [cell.strip() for cell in line.strip().strip("|").split("|")]


def _is_table_separator_row(cells: List[str]) -> bool:
    if not cells:
        return False
    for cell in cells:
        normalized = cell.replace(":", "").replace("-", "").strip()
        if normalized:
            return False
    return True


def _parse_bool(value: str) -> bool:
    return value.strip().lower() in {"1", "true", "yes", "y", "on"}


def _strip_wrapping_quotes(value: str) -> str:
    value = value.strip()
    if len(value) >= 2 and value[0] == value[-1] and value[0] in {"'", '"'}:
        return value[1:-1]
    return value


def _normalize_replace_mode(mode: str) -> str:
    normalized = (mode or "auto").strip().lower().replace("-", "_")
    aliases = {
        "exact": "exact_restore",
        "restore": "exact_restore",
        "editable": "editable_rebuild",
        "rebuild": "editable_rebuild",
        "template": "editable_rebuild_with_template",
        "editable_with_template": "editable_rebuild_with_template",
        "rebuild_with_template": "editable_rebuild_with_template",
    }
    normalized = aliases.get(normalized, normalized)
    if normalized not in REPLACE_MODES:
        allowed = ", ".join(sorted(REPLACE_MODES))
        raise ValueError(f"Unsupported markdown replace mode '{mode}'. Supported modes: {allowed}.")
    return normalized


def _match_leading_directive(pattern: re.Pattern, markdown_text: str):
    normalized_text = markdown_text.lstrip("\ufeff \t\r\n")
    offset = len(markdown_text) - len(normalized_text)
    return pattern.match(normalized_text), offset


def _parse_toc_directive(comment_text: str) -> Optional[Dict[str, Any]]:
    """Parse a Markdown TOC directive into a native Word TOC block."""
    match = TOC_RE.match(comment_text.strip())
    if not match:
        return None

    body = match.group("body").strip()
    block: Dict[str, Any] = {
        "type": "toc",
        "title": "Contents",
        "max_level": 3,
        "add_page_break_after": False,
        "toc_style": "dotted",
    }
    if not body:
        return block

    metadata: Dict[str, str] = {}
    for raw_line in body.splitlines():
        line = raw_line.strip()
        if not line or line.startswith("<!--") or line == "-->":
            continue
        if ":" in line:
            key, value = line.split(":", 1)
            metadata[key.strip().lower().replace("-", "_")] = _strip_wrapping_quotes(value)

    for key, value in re.findall(r"([A-Za-z_][\w-]*)\s*=\s*(\"[^\"]*\"|'[^']*'|[^\s]+)", body):
        metadata[key.strip().lower().replace("-", "_")] = _strip_wrapping_quotes(value)

    if "title" in metadata:
        block["title"] = metadata["title"]
    if metadata.get("title", "").strip().lower() in {"none", "false", "off"}:
        block["title"] = ""
    if "max_level" in metadata:
        try:
            block["max_level"] = max(1, min(int(metadata["max_level"]), 9))
        except ValueError:
            pass
    if "level" in metadata and "max_level" not in metadata:
        try:
            block["max_level"] = max(1, min(int(metadata["level"]), 9))
        except ValueError:
            pass
    if "add_page_break_after" in metadata:
        block["add_page_break_after"] = _parse_bool(metadata["add_page_break_after"])
    if "page_break_after" in metadata:
        block["add_page_break_after"] = _parse_bool(metadata["page_break_after"])
    if "style" in metadata:
        block["toc_style"] = metadata["style"]
    if "toc_style" in metadata:
        block["toc_style"] = metadata["toc_style"]
    return block


def _normalize_markdown_image_path(path_text: str) -> str:
    """Normalize markdown image path syntax."""
    path_text = path_text.strip()
    if path_text.startswith("<") and path_text.endswith(">"):
        path_text = path_text[1:-1].strip()
    return path_text


def _resolve_markdown_image_path(path_text: str, base_dir: Optional[str] = None) -> str:
    """Resolve markdown image paths for local document insertion."""
    candidate = os.path.expanduser(_normalize_markdown_image_path(path_text))

    if os.name == "nt":
        if re.match(r"^/[A-Za-z]:/", candidate):
            candidate = candidate[1:]
        elif candidate.startswith("/Users/"):
            drive = os.path.splitdrive(os.path.expanduser("~"))[0] or "C:"
            candidate = drive + candidate.replace("/", os.sep)

    if os.path.isabs(candidate):
        return os.path.normpath(candidate)
    if base_dir:
        return os.path.normpath(os.path.join(base_dir, candidate))
    return os.path.normpath(candidate)


def _extract_image_to_media_dir(doc, rel_id: str, media_dir: str) -> Optional[str]:
    """Extract an embedded image to a sibling media directory and return the absolute path."""
    image_part = doc.part.related_parts.get(rel_id)
    if image_part is None:
        return None

    os.makedirs(media_dir, exist_ok=True)
    image_name = os.path.basename(str(image_part.partname))
    output_path = os.path.abspath(os.path.join(media_dir, image_name))
    if not os.path.exists(output_path):
        with open(output_path, "wb") as image_file:
            image_file.write(image_part.blob)
    return output_path


def _get_paragraph_image_blocks(doc, para, media_dir: str) -> List[Dict[str, str]]:
    """Extract markdown image blocks from a paragraph's drawing elements."""
    blocks: List[Dict[str, str]] = []
    for node in para._element.iter():
        if node.tag != qn("w:drawing"):
            continue

        rel_id = None
        alt_text = ""
        for child in node.iter():
            if child.tag == qn("a:blip"):
                rel_id = child.get(qn("r:embed"))
            elif child.tag == qn("wp:docPr"):
                alt_text = (
                    child.get("descr")
                    or child.get("title")
                    or child.get("name")
                    or ""
                ).strip()

        if not rel_id:
            continue

        extracted_path = _extract_image_to_media_dir(doc, rel_id, media_dir)
        if not extracted_path:
            continue

        blocks.append(
            {
                "type": "image",
                "alt": alt_text or os.path.basename(extracted_path),
                "path": extracted_path,
            }
        )
    return blocks


def _paragraph_alignment_name(para) -> Optional[str]:
    alignment_map = {
        WD_ALIGN_PARAGRAPH.LEFT: "left",
        WD_ALIGN_PARAGRAPH.CENTER: "center",
        WD_ALIGN_PARAGRAPH.RIGHT: "right",
        WD_ALIGN_PARAGRAPH.JUSTIFY: "justify",
    }
    return alignment_map.get(para.alignment)


def _paragraph_has_page_break(para) -> bool:
    for node in para._element.iter():
        if node.tag == qn("w:br") and node.get(qn("w:type")) == "page":
            return True
    return False


def _paragraph_has_section_break(para) -> bool:
    p_pr = para._element.find(qn("w:pPr"))
    return p_pr is not None and p_pr.find(qn("w:sectPr")) is not None


def _paragraph_exports_page_boundary(para) -> bool:
    return _paragraph_has_page_break(para) or _paragraph_has_section_break(para)


def _append_exported_boundaries(blocks: List[Dict[str, Any]], para) -> None:
    if _paragraph_has_page_break(para):
        blocks.append({"type": "page_break"})
    if _paragraph_has_section_break(para):
        blocks.append({"type": "section_break"})


def _paragraph_field_instructions(para) -> List[str]:
    instructions: List[str] = []
    for instr_text in para._element.iter(qn("w:instrText")):
        if instr_text.text:
            instructions.append(instr_text.text.strip())
    return instructions


def _toc_block_from_paragraph(para) -> Optional[Dict[str, Any]]:
    for instruction in _paragraph_field_instructions(para):
        normalized = " ".join(instruction.split())
        if not normalized.upper().startswith("TOC "):
            continue

        max_level = 3
        level_match = re.search(r'\\o\s+"1-(\d+)"', normalized)
        if level_match:
            max_level = max(1, min(int(level_match.group(1)), 9))

        toc_style = "dotted"
        if re.search(r'\\n\s+"1-\d+"', normalized):
            toc_style = "links"
        elif re.search(r'\\p\s+" "', normalized):
            toc_style = "page_numbers"

        return {
            "type": "toc",
            "title": "",
            "max_level": max_level,
            "toc_style": toc_style,
            "add_page_break_after": False,
        }
    return None


def _build_fidelity_bundle(doc_path: str) -> str:
    """Encode the original DOCX as a markdown comment for exact round-tripping."""
    with open(doc_path, "rb") as docx_file:
        payload = docx_file.read()

    encoded = base64.b64encode(payload).decode("ascii")
    wrapped = "\n".join(textwrap.wrap(encoded, width=120))
    digest = hashlib.sha256(payload).hexdigest()
    filename = os.path.basename(doc_path)

    return (
        "<!-- wadocx:fidelity-bundle\n"
        f"version: {FIDELITY_VERSION}\n"
        f"filename: {filename}\n"
        f"sha256: {digest}\n"
        "encoding: base64\n"
        "data:\n"
        f"{wrapped}\n"
        "-->"
    )


def _extract_fidelity_bundle(markdown_text: str) -> Optional[Dict[str, Any]]:
    """Parse a fidelity bundle comment from markdown text."""
    match, offset = _match_leading_directive(FIDELITY_RE, markdown_text)
    if not match:
        return None

    lines = match.group("body").splitlines()
    metadata: Dict[str, str] = {}
    data_lines: List[str] = []
    in_data = False

    for raw_line in lines:
        line = raw_line.rstrip()
        if not in_data:
            if line.strip() == "data:":
                in_data = True
                continue
            if ":" in line:
                key, value = line.split(":", 1)
                metadata[key.strip().lower()] = value.strip()
            continue
        if line.strip():
            data_lines.append(line.strip())

    if not data_lines:
        return None

    payload = base64.b64decode("".join(data_lines))
    return {
        "metadata": metadata,
        "payload": payload,
        "match_end": offset + match.end(),
    }


def _restore_docx_from_fidelity_bundle(doc_path: str, bundle: Dict[str, Any]) -> Dict[str, Any]:
    """Write the exact DOCX bytes from a fidelity bundle to disk."""
    payload = bundle["payload"]
    os.makedirs(os.path.dirname(os.path.abspath(doc_path)), exist_ok=True)
    with open(doc_path, "wb") as docx_file:
        docx_file.write(payload)

    return {
        "inserted_blocks": 0,
        "block_count": 0,
        "restored_exact_docx": True,
        "sha256": hashlib.sha256(payload).hexdigest(),
    }


def _extract_base_template_directive(
    markdown_text: str,
    source_base_dir: Optional[str] = None,
) -> Optional[Dict[str, Any]]:
    """Parse an optional base-template directive from markdown text."""
    match, offset = _match_leading_directive(BASE_TEMPLATE_RE, markdown_text)
    if not match:
        return None

    metadata: Dict[str, str] = {}
    for raw_line in match.group("body").splitlines():
        line = raw_line.strip()
        if not line or ":" not in line:
            continue
        key, value = line.split(":", 1)
        metadata[key.strip().lower()] = value.strip()

    template_path = _strip_wrapping_quotes(metadata.get("path", ""))
    if not template_path:
        return None
    template_path = os.path.expanduser(template_path)
    if source_base_dir and not os.path.isabs(template_path):
        template_path = os.path.join(source_base_dir, template_path)

    return {
        "metadata": metadata,
        "template_path": os.path.normpath(template_path),
        "match_end": offset + match.end(),
    }


def _extract_template_section_breaks(doc) -> List[Any]:
    """Collect intermediate section-break paragraphs from a template document."""
    section_breaks: List[Any] = []
    for el in get_body_elements(doc):
        if not is_paragraph_element(el):
            continue
        p_pr = el.find(qn("w:pPr"))
        sect_pr = p_pr.find(qn("w:sectPr")) if p_pr is not None else None
        if sect_pr is not None:
            section_breaks.append(deepcopy(el))
    return section_breaks


def parse_markdown_blocks(markdown_text: str) -> List[Dict[str, Any]]:
    """Parse a markdown string into simple document blocks."""
    blocks: List[Dict[str, Any]] = []
    paragraph_lines: List[str] = []
    code_lines: Optional[List[str]] = None
    current_alignment: Optional[str] = None
    lines = markdown_text.splitlines()
    i = 0

    def flush_paragraph() -> None:
        nonlocal paragraph_lines
        if paragraph_lines:
            text = " ".join(line.strip() for line in paragraph_lines if line.strip())
            text = _normalize_inline_markdown_text(text)
            if text:
                block = {"type": "paragraph", "text": text}
                if current_alignment:
                    block["alignment"] = current_alignment
                blocks.append(block)
            paragraph_lines = []

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if stripped.startswith("```"):
            if code_lines is None:
                flush_paragraph()
                code_lines = []
            else:
                blocks.append({"type": "paragraph", "text": "\n".join(code_lines)})
                code_lines = None
            i += 1
            continue

        if code_lines is not None:
            code_lines.append(line)
            i += 1
            continue

        if not stripped:
            flush_paragraph()
            i += 1
            continue

        div_open_match = DIV_OPEN_RE.match(stripped)
        if div_open_match:
            flush_paragraph()
            current_alignment = div_open_match.group("align").lower()
            i += 1
            continue

        if DIV_CLOSE_RE.match(stripped):
            flush_paragraph()
            current_alignment = None
            i += 1
            continue

        if PAGE_BREAK_RE.match(stripped):
            flush_paragraph()
            blocks.append({"type": "page_break"})
            i += 1
            continue

        if SECTION_BREAK_RE.match(stripped):
            flush_paragraph()
            blocks.append({"type": "section_break"})
            i += 1
            continue

        lowered = stripped.lower()
        if lowered.startswith("<!--toc") or lowered.startswith("<!-- toc") or lowered.startswith("<!-- wadocx:toc"):
            flush_paragraph()
            comment_lines = [line]
            while "-->" not in comment_lines[-1] and i + 1 < len(lines):
                i += 1
                comment_lines.append(lines[i])
            toc_block = _parse_toc_directive("\n".join(comment_lines))
            if toc_block:
                if current_alignment:
                    toc_block["alignment"] = current_alignment
                blocks.append(toc_block)
                i += 1
                continue
            paragraph_lines.extend(comment_lines)
            i += 1
            continue

        heading_match = HEADING_RE.match(stripped)
        if heading_match:
            flush_paragraph()
            blocks.append(
                {
                    "type": "heading",
                    "level": len(heading_match.group(1)),
                    "text": _normalize_inline_markdown_text(heading_match.group(2).strip()),
                    **({"alignment": current_alignment} if current_alignment else {}),
                }
            )
            i += 1
            continue

        if _is_markdown_table_line(stripped):
            flush_paragraph()
            table_lines = []
            while i < len(lines) and _is_markdown_table_line(lines[i].strip()):
                table_lines.append(lines[i].strip())
                i += 1

            rows = [_split_markdown_row(table_line) for table_line in table_lines]
            if len(rows) > 1 and _is_table_separator_row(rows[1]):
                rows.pop(1)
            if rows:
                blocks.append({"type": "table", "rows": rows})
            continue

        ordered_match = ORDERED_LIST_RE.match(line)
        bullet_match = BULLET_LIST_RE.match(line)
        if ordered_match or bullet_match:
            flush_paragraph()
            ordered = ordered_match is not None
            items = []
            while i < len(lines):
                current_line = lines[i]
                current_match = ORDERED_LIST_RE.match(current_line) if ordered else BULLET_LIST_RE.match(current_line)
                if not current_match:
                    break
                items.append(_normalize_inline_markdown_text(current_match.group(2 if ordered else 1).strip()))
                i += 1
            block = {"type": "list", "ordered": ordered, "items": items}
            if current_alignment:
                block["alignment"] = current_alignment
            blocks.append(block)
            continue

        image_match = IMAGE_RE.match(stripped)
        if image_match:
            flush_paragraph()
            block = {
                "type": "image",
                "alt": image_match.group("alt").strip(),
                "path": _normalize_markdown_image_path(image_match.group("path")),
            }
            if current_alignment:
                block["alignment"] = current_alignment
            blocks.append(block)
            i += 1
            continue

        paragraph_lines.append(line)
        i += 1

    flush_paragraph()
    if code_lines:
        blocks.append({"type": "paragraph", "text": "\n".join(code_lines)})
    return blocks


def _paragraph_list_kind(para) -> Optional[str]:
    """Infer whether a paragraph belongs to a bullet or numbered list."""
    if para.style:
        style_name = para.style.name.lower()
        if "list number" in style_name or style_name == "numbered list":
            return "ordered"
        if "list bullet" in style_name or style_name == "bullet list":
            return "unordered"

    p_pr = para._element.find(qn("w:pPr"))
    if p_pr is None:
        return None
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is None:
        return None
    num_id = num_pr.find(qn("w:numId"))
    if num_id is not None and num_id.get(qn("w:val")) == "2":
        return "ordered"
    return "unordered"


def document_to_markdown(doc_path: str, include_fidelity_bundle: bool = True) -> str:
    """Export a Word document to a simple markdown representation."""
    if not os.path.exists(doc_path):
        return f"Document {doc_path} does not exist"

    doc = Document(doc_path)
    blocks: List[Dict[str, Any]] = []
    media_dir = os.path.splitext(os.path.abspath(doc_path))[0] + "_media"

    for el in get_body_elements(doc):
        if is_paragraph_element(el):
            para = get_paragraph_from_element(doc, el)
            if para is None:
                continue
            toc_block = _toc_block_from_paragraph(para)
            if toc_block:
                if blocks and blocks[-1].get("toc_title"):
                    title_block = blocks.pop()
                    toc_block["title"] = title_block["text"]
                    if title_block.get("alignment"):
                        toc_block["alignment"] = title_block["alignment"]
                blocks.append(toc_block)
                _append_exported_boundaries(blocks, para)
                continue

            image_blocks = _get_paragraph_image_blocks(doc, para, media_dir)
            text = para.text.rstrip()
            alignment = _paragraph_alignment_name(para)
            if alignment:
                for image_block in image_blocks:
                    image_block["alignment"] = alignment
            if not text:
                blocks.extend(image_blocks)
                _append_exported_boundaries(blocks, para)
                continue

            style_name = para.style.name if para.style else ""
            if style_name.startswith("Heading "):
                try:
                    level = int(style_name.split(" ")[1])
                except (IndexError, ValueError):
                    level = 1
                block = {"type": "heading", "level": level, "text": text}
                if alignment:
                    block["alignment"] = alignment
                blocks.append(block)
                blocks.extend(image_blocks)
                _append_exported_boundaries(blocks, para)
                continue

            if style_name == "TOC Heading":
                block = {"type": "paragraph", "text": text, "toc_title": True}
                if alignment:
                    block["alignment"] = alignment
                blocks.append(block)
                blocks.extend(image_blocks)
                _append_exported_boundaries(blocks, para)
                continue

            list_kind = _paragraph_list_kind(para)
            if list_kind:
                ordered = list_kind == "ordered"
                if blocks and blocks[-1]["type"] == "list" and blocks[-1]["ordered"] == ordered:
                    blocks[-1]["items"].append(text)
                else:
                    block = {"type": "list", "ordered": ordered, "items": [text]}
                    if alignment:
                        block["alignment"] = alignment
                    blocks.append(block)
                blocks.extend(image_blocks)
                _append_exported_boundaries(blocks, para)
                continue

            block = {"type": "paragraph", "text": text}
            if alignment:
                block["alignment"] = alignment
            blocks.append(block)
            blocks.extend(image_blocks)
            _append_exported_boundaries(blocks, para)
            continue

        if is_table_element(el):
            table = Table(el, doc._body)
            rows = []
            for row in table.rows:
                rows.append([normalize_paragraph_text(cell.text) for cell in row.cells])
            if rows:
                blocks.append({"type": "table", "rows": rows})

    rendered_blocks: List[str] = []

    def render_aligned(markdown_block: str, block: Dict[str, Any]) -> str:
        alignment = block.get("alignment")
        if not alignment:
            return markdown_block
        return f'<div align="{alignment}">\n{markdown_block}\n</div>'

    for block in blocks:
        if block["type"] == "heading":
            rendered_blocks.append(render_aligned(f"{'#' * block['level']} {block['text']}", block))
            continue

        if block["type"] == "list":
            prefix = "1." if block["ordered"] else "-"
            lines = []
            for index, item in enumerate(block["items"], start=1):
                marker = f"{index}." if block["ordered"] else prefix
                lines.append(f"{marker} {item}")
            rendered_blocks.append(render_aligned("\n".join(lines), block))
            continue

        if block["type"] == "table":
            rows = block["rows"]
            if not rows:
                continue
            width = max(len(row) for row in rows)

            def pad(row: List[str]) -> List[str]:
                return row + [""] * (width - len(row))

            header = pad(rows[0])
            lines = [
                f"| {' | '.join(header)} |",
                f"| {' | '.join(['---'] * width)} |",
            ]
            for row in rows[1:]:
                lines.append(f"| {' | '.join(pad(row))} |")
            rendered_blocks.append(render_aligned("\n".join(lines), block))
            continue

        if block["type"] == "image":
            alt_text = block.get("alt", "").replace("]", "\\]")
            image_path = block.get("path", "")
            rendered_blocks.append(render_aligned(f"![{alt_text}](<{image_path}>)", block))
            continue

        if block["type"] == "page_break":
            rendered_blocks.append("<!-- PAGE BREAK -->")
            continue

        if block["type"] == "section_break":
            rendered_blocks.append("<!-- SECTION BREAK -->")
            continue

        if block["type"] == "toc":
            title = block.get("title", "Contents")
            title_value = title if title else "none"
            max_level = block.get("max_level", 3)
            toc_style = block.get("toc_style", "dotted")
            rendered_blocks.append(render_aligned(
                "<!-- wadocx:toc\n"
                f"title: {title_value}\n"
                f"max_level: {max_level}\n"
                f"style: {toc_style}\n"
                "-->",
                block,
            ))
            continue

        rendered_blocks.append(render_aligned(block["text"], block))

    markdown_body = "\n\n".join(rendered_blocks)
    if not include_fidelity_bundle:
        return markdown_body
    fidelity_bundle = _build_fidelity_bundle(doc_path)
    return f"{fidelity_bundle}\n\n{markdown_body}"


def export_document_markdown(
    doc_path: str,
    output_path: Optional[str] = None,
    include_fidelity_bundle: bool = True,
) -> str:
    """Export a document to markdown on disk and return the output path."""
    markdown = document_to_markdown(doc_path, include_fidelity_bundle=include_fidelity_bundle)
    if markdown.startswith("Document ") and markdown.endswith(" does not exist"):
        return markdown

    if output_path is None:
        base_name, _ = os.path.splitext(doc_path)
        output_path = f"{base_name}.md"

    markdown_dir = os.path.dirname(os.path.abspath(output_path)) or os.getcwd()
    output_stem = os.path.splitext(os.path.basename(output_path))[0]
    output_media_dir = os.path.join(markdown_dir, f"{output_stem}_media")
    media_copies: Dict[str, str] = {}

    def relativize_image_link(match: re.Match) -> str:
        image_path = match.group("path")
        if not os.path.isabs(image_path):
            return match.group(0)
        target_path = media_copies.get(image_path)
        if target_path is None:
            os.makedirs(output_media_dir, exist_ok=True)
            target_path = os.path.join(output_media_dir, os.path.basename(image_path))
            if os.path.abspath(image_path) != os.path.abspath(target_path):
                shutil.copy2(image_path, target_path)
            media_copies[image_path] = target_path
        try:
            relative_path = os.path.relpath(target_path, markdown_dir)
        except ValueError:
            return match.group(0)
        relative_path = relative_path.replace(os.sep, "/")
        return f"![{match.group('alt')}](<{relative_path}>)"

    markdown = IMAGE_LINK_RE.sub(relativize_image_link, markdown)

    with open(output_path, "w", encoding="utf-8") as markdown_file:
        markdown_file.write(markdown)

    return output_path


def clear_document_body(doc) -> None:
    """Remove all body elements except section properties."""
    body = doc.element.body
    for child in list(body.iterchildren()):
        if child.tag == qn("w:sectPr"):
            continue
        body.remove(child)


def _insert_page_break_after_element(
    doc,
    anchor_element,
):
    """Insert a simple page break after the anchor."""
    new_para = doc.add_paragraph("")
    new_para.add_run().add_break(WD_BREAK.PAGE)
    anchor_element.addnext(new_para._element)
    return new_para._element, 1


def _insert_section_break_after_element(
    doc,
    anchor_element,
    section_break_elements: Optional[List[Any]] = None,
):
    """Insert a template section break when available, otherwise fall back to a page break."""
    if section_break_elements:
        section_break = section_break_elements.pop(0)
        anchor_element.addnext(section_break)
        return section_break, 1
    return _insert_page_break_after_element(doc, anchor_element)


def replace_document_with_markdown(
    doc_path: str,
    markdown_text: str,
    mode: str = "auto",
    source_base_dir: Optional[str] = None,
) -> Dict[str, Any]:
    """Replace the body of a document with parsed markdown blocks."""
    replace_mode = _normalize_replace_mode(mode)
    fidelity_bundle = _extract_fidelity_bundle(markdown_text)
    if fidelity_bundle and replace_mode in {"auto", "exact_restore"}:
        return _restore_docx_from_fidelity_bundle(doc_path, fidelity_bundle)
    if replace_mode == "exact_restore":
        return {"error": "Exact restore mode requires a wadocx:fidelity-bundle."}

    section_break_elements: List[Any] = []
    if fidelity_bundle and replace_mode == "editable_rebuild_with_template":
        _restore_docx_from_fidelity_bundle(doc_path, fidelity_bundle)
        template_doc = Document(doc_path)
        section_break_elements = _extract_template_section_breaks(template_doc)
        markdown_text = markdown_text[fidelity_bundle["match_end"] :].lstrip()
    elif fidelity_bundle and replace_mode == "editable_rebuild":
        markdown_text = markdown_text[fidelity_bundle["match_end"] :].lstrip()

    base_template = _extract_base_template_directive(markdown_text, source_base_dir)
    if base_template:
        template_path = base_template["template_path"]
        if not os.path.exists(template_path):
            return {"error": f"Base template markdown does not exist: {template_path}"}
        with open(template_path, "r", encoding="utf-8") as template_file:
            template_markdown = template_file.read()
        template_bundle = _extract_fidelity_bundle(template_markdown)
        if not template_bundle:
            return {
                "error": (
                    "Base template markdown must include a wadocx:fidelity-bundle "
                    f"to preserve template structure: {template_path}"
                )
            }
        _restore_docx_from_fidelity_bundle(doc_path, template_bundle)
        template_doc = Document(doc_path)
        section_break_elements = _extract_template_section_breaks(template_doc)
        markdown_text = markdown_text[base_template["match_end"] :].lstrip()

    doc = Document(doc_path) if os.path.exists(doc_path) else Document()
    clear_document_body(doc)

    blocks = parse_markdown_blocks(markdown_text)
    base_dir = source_base_dir or os.path.dirname(os.path.abspath(doc_path))
    for block in blocks:
        if block.get("type") == "image":
            block["path"] = _resolve_markdown_image_path(block.get("path", ""), base_dir)
    anchor = doc.add_paragraph("")
    inserted = 0
    current_element = anchor._element
    for block in blocks:
        if block.get("type") == "page_break":
            current_element, added = _insert_page_break_after_element(
                doc,
                current_element,
            )
            inserted += added
            continue
        if block.get("type") == "section_break":
            current_element, added = _insert_section_break_after_element(
                doc,
                current_element,
                section_break_elements,
            )
            inserted += added
            continue
        added = insert_content_blocks_after_element(doc, current_element, [block])
        if added:
            new_current = current_element
            for _ in range(added):
                new_current = new_current.getnext()
                if new_current is None:
                    break
            if new_current is not None:
                current_element = new_current
            else:
                current_element = current_element.getnext()
            inserted += added
    anchor._element.getparent().remove(anchor._element)
    doc.save(doc_path)

    return {"inserted_blocks": inserted, "block_count": len(blocks)}


def replace_section_with_markdown(doc_path: str, header_text: str, markdown_text: str) -> Dict[str, Any]:
    """Replace a section body with parsed markdown blocks."""
    if not os.path.exists(doc_path):
        return {"error": f"Document {doc_path} does not exist"}

    fidelity_bundle = _extract_fidelity_bundle(markdown_text)
    if fidelity_bundle:
        return {
            "error": (
                "Section replacement does not accept a wadocx fidelity bundle. "
                "Use replace_document_with_markdown for exact document restoration."
            )
        }

    doc = Document(doc_path)
    header_el, removed_count = delete_block_under_header(doc, header_text)
    if header_el is None:
        return {"error": f"Header '{header_text}' not found in document."}

    blocks = parse_markdown_blocks(markdown_text)
    base_dir = os.path.dirname(os.path.abspath(doc_path))
    for block in blocks:
        if block.get("type") == "image":
            block["path"] = _resolve_markdown_image_path(block.get("path", ""), base_dir)
    inserted = insert_content_blocks_after_element(doc, header_el, blocks)
    doc.save(doc_path)
    return {"inserted_blocks": inserted, "removed_blocks": removed_count, "block_count": len(blocks)}

