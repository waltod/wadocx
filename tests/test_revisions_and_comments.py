import asyncio
import json
import zipfile
from pathlib import Path

from docx import Document
from lxml import etree

from word_document_server.tools import revision_tools as RT
from word_document_server.tools import comment_tools as CT


def _make_doc(path: Path) -> None:
    doc = Document()
    doc.add_paragraph("The quick brown fox jumps over the lazy dog.")
    doc.add_paragraph("Hello world, this is a second paragraph.")
    doc.save(str(path))


# --- tracked changes ---------------------------------------------------------

def test_enable_track_changes(tmp_path: Path):
    path = tmp_path / "tc.docx"
    _make_doc(path)
    assert "enabled" in asyncio.run(RT.enable_track_changes(str(path), True))
    with zipfile.ZipFile(path) as zf:
        settings = zf.read("word/settings.xml").decode()
    assert "trackChanges" in settings


def test_tracked_replace_accept(tmp_path: Path):
    path = tmp_path / "tr_accept.docx"
    _make_doc(path)
    msg = asyncio.run(RT.tracked_search_and_replace(str(path), "quick", "slow", author="Rev"))
    assert "Tracked-replaced 1" in msg

    counts = json.loads(asyncio.run(RT.get_revision_counts(str(path))))
    assert counts["insertions"] == 1
    assert counts["deletions"] == 1

    # The change is marked with the author.
    with zipfile.ZipFile(path) as zf:
        doc_xml = zf.read("word/document.xml").decode()
    assert 'w:author="Rev"' in doc_xml
    assert "<w:del" in doc_xml and "<w:ins" in doc_xml

    asyncio.run(RT.accept_all_changes(str(path)))
    assert Document(str(path)).paragraphs[0].text == "The slow brown fox jumps over the lazy dog."


def test_tracked_replace_reject(tmp_path: Path):
    path = tmp_path / "tr_reject.docx"
    _make_doc(path)
    asyncio.run(RT.tracked_search_and_replace(str(path), "quick", "slow", author="Rev"))
    asyncio.run(RT.reject_all_changes(str(path)))
    assert Document(str(path)).paragraphs[0].text == "The quick brown fox jumps over the lazy dog."


def test_mark_text_as_deleted_roundtrip(tmp_path: Path):
    path = tmp_path / "del.docx"
    _make_doc(path)
    asyncio.run(RT.mark_text_as_deleted(str(path), "lazy ", author="Rev"))
    # python-docx .text ignores deleted text, so it disappears from .text.
    assert "lazy" not in Document(str(path)).paragraphs[0].text
    asyncio.run(RT.reject_all_changes(str(path)))
    assert "lazy" in Document(str(path)).paragraphs[0].text


def test_add_tracked_insertion(tmp_path: Path):
    path = tmp_path / "ins.docx"
    _make_doc(path)
    asyncio.run(RT.add_tracked_insertion(str(path), "A new tracked line.", author="Rev"))
    with zipfile.ZipFile(path) as zf:
        doc_xml = zf.read("word/document.xml").decode()
    assert "<w:ins" in doc_xml
    # Accepting keeps the inserted line.
    asyncio.run(RT.accept_all_changes(str(path)))
    assert any("A new tracked line." == p.text for p in Document(str(path)).paragraphs)


# --- comments ----------------------------------------------------------------

def test_add_comment(tmp_path: Path):
    path = tmp_path / "comment.docx"
    _make_doc(path)
    result = json.loads(asyncio.run(
        CT.add_comment(str(path), "Please review this.", search_text="brown fox", author="Alice", initials="AL")
    ))
    assert result["success"] is True
    # Document still reloads and the comment is present.
    out = json.loads(asyncio.run(CT.get_all_comments(str(path))))
    assert out["total_comments"] >= 1


def test_reply_and_resolve_comment(tmp_path: Path):
    path = tmp_path / "thread.docx"
    _make_doc(path)
    created = json.loads(asyncio.run(
        CT.add_comment(str(path), "Top-level note.", search_text="quick brown", author="Alice", initials="AL")
    ))
    cid = created["comment_id"]

    reply = json.loads(asyncio.run(
        CT.reply_to_comment(str(path), cid, "Replying to the note.", author="Bob", initials="BO")
    ))
    assert reply["success"] is True

    resolved = json.loads(asyncio.run(CT.resolve_comment(str(path), cid, True)))
    assert resolved["success"] is True

    # commentsExtended is present and well-formed; both comments survive a reload.
    with zipfile.ZipFile(path) as zf:
        names = zf.namelist()
        assert "word/commentsExtended.xml" in names
        etree.fromstring(zf.read("word/commentsExtended.xml"))
        etree.fromstring(zf.read("word/comments.xml"))
    out = json.loads(asyncio.run(CT.get_all_comments(str(path))))
    assert out["total_comments"] >= 2


def test_add_comment_text_not_found(tmp_path: Path):
    path = tmp_path / "nf.docx"
    _make_doc(path)
    result = json.loads(asyncio.run(
        CT.add_comment(str(path), "note", search_text="nonexistent phrase")
    ))
    assert result["success"] is False
