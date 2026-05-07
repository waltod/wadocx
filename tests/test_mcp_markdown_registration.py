import asyncio
import base64

from docx import Document

from word_document_server import main


def _tool_text(result):
    return "\n".join(part.text for part in result.content if hasattr(part, "text"))


def _write_png(path):
    png_base64 = (
        "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mP8/x8AAwMCAO7Z0mQAAAAASUVORK5CYII="
    )
    path.write_bytes(base64.b64decode(png_base64))


def test_mcp_registers_one_call_markdown_tools_with_expected_parameters():
    main.register_tools()
    tools = {tool.name: tool for tool in asyncio.run(main.mcp.list_tools())}

    assert "get_document_markdown" in tools
    assert "export_document_markdown" in tools
    assert "replace_document_with_markdown" in tools
    assert "replace_document_with_markdown_file" in tools

    assert "include_fidelity_bundle" in tools["get_document_markdown"].parameters["properties"]
    assert "include_fidelity_bundle" in tools["export_document_markdown"].parameters["properties"]
    assert "mode" in tools["replace_document_with_markdown"].parameters["properties"]
    assert "source_base_dir" in tools["replace_document_with_markdown"].parameters["properties"]
    assert "mode" in tools["replace_document_with_markdown_file"].parameters["properties"]
    assert "markdown_path" in tools["replace_document_with_markdown_file"].parameters["properties"]
    assert tools["get_document_markdown"].annotations.readOnlyHint in {None, False}
    assert tools["replace_document_with_markdown"].parameters["properties"]["mode"]["enum"] == [
        "auto",
        "exact_restore",
        "editable_rebuild",
        "editable_rebuild_with_template",
    ]
    assert tools["replace_document_with_markdown_file"].parameters["properties"]["mode"]["enum"] == [
        "auto",
        "exact_restore",
        "editable_rebuild",
        "editable_rebuild_with_template",
    ]


def test_mcp_markdown_tools_export_and_import_through_call_tool(tmp_path):
    main.register_tools()
    source_doc_path = tmp_path / "source.docx"
    exported_md_path = tmp_path / "source.md"
    restored_doc_path = tmp_path / "restored.docx"

    doc = Document()
    doc.add_heading("MCP Round Trip", level=1)
    doc.add_paragraph("Body through MCP call_tool.")
    doc.save(source_doc_path)

    export_result = asyncio.run(
        main.mcp.call_tool(
            "export_document_markdown",
            {
                "filename": str(source_doc_path),
                "output_filename": str(exported_md_path),
            },
        )
    )
    import_result = asyncio.run(
        main.mcp.call_tool(
            "replace_document_with_markdown_file",
            {
                "filename": str(restored_doc_path),
                "markdown_path": str(exported_md_path),
            },
        )
    )
    invalid_mode_result = asyncio.run(
        main.mcp.call_tool(
            "replace_document_with_markdown_file",
            {
                "filename": str(restored_doc_path),
                "markdown_path": str(exported_md_path),
                "mode": "bad-mode",
            },
        )
    )

    assert "Markdown exported to" in _tool_text(export_result)
    assert exported_md_path.exists()
    assert "restored exactly from a wadocx fidelity bundle" in _tool_text(import_result)
    assert source_doc_path.read_bytes() == restored_doc_path.read_bytes()
    assert "Unsupported markdown replace mode" in _tool_text(invalid_mode_result)


def test_mcp_markdown_text_import_and_export_errors_through_call_tool(tmp_path):
    main.register_tools()
    text_doc_path = tmp_path / "text-import.docx"
    source_doc_path = tmp_path / "source.docx"
    missing_output_path = tmp_path / "missing" / "out.md"

    source = Document()
    source.add_heading("Source", level=1)
    source.save(source_doc_path)

    import_result = asyncio.run(
        main.mcp.call_tool(
            "replace_document_with_markdown",
            {
                "filename": str(text_doc_path),
                "markdown_text": "# MCP Text Import\n\nBody through text input.",
            },
        )
    )
    invalid_mode_result = asyncio.run(
        main.mcp.call_tool(
            "replace_document_with_markdown",
            {
                "filename": str(text_doc_path),
                "markdown_text": "# MCP Text Import\n",
                "mode": "bad-mode",
            },
        )
    )
    export_error_result = asyncio.run(
        main.mcp.call_tool(
            "export_document_markdown",
            {
                "filename": str(source_doc_path),
                "output_filename": str(missing_output_path),
            },
        )
    )

    assert "replaced from markdown" in _tool_text(import_result)
    assert "Body through text input." in [p.text for p in Document(text_doc_path).paragraphs]
    assert "Unsupported markdown replace mode" in _tool_text(invalid_mode_result)
    assert "Markdown export failed:" in _tool_text(export_error_result)


def test_mcp_markdown_file_import_resolves_relative_paths_for_editable_rebuild(tmp_path):
    main.register_tools()
    source_doc_path = tmp_path / "source.docx"
    editable_md_path = tmp_path / "editable.md"
    draft_md_path = tmp_path / "draft.md"
    image_path = tmp_path / "risk-flow.png"
    rebuilt_doc_path = tmp_path / "rebuilt.docx"
    _write_png(image_path)

    doc = Document()
    doc.add_heading("Editable MCP Export", level=1)
    doc.add_paragraph("Body through editable MCP export.")
    doc.save(source_doc_path)

    export_result = asyncio.run(
        main.mcp.call_tool(
            "export_document_markdown",
            {
                "filename": str(source_doc_path),
                "output_filename": str(editable_md_path),
                "include_fidelity_bundle": False,
            },
        )
    )
    draft_md_path.write_text(
        """# Editable MCP Import

![Risk flow](<risk-flow.png>)

Body with a relative image.
""",
        encoding="utf-8",
    )
    import_result = asyncio.run(
        main.mcp.call_tool(
            "replace_document_with_markdown_file",
            {
                "filename": str(rebuilt_doc_path),
                "markdown_path": str(draft_md_path),
                "mode": "editable_rebuild",
            },
        )
    )

    assert "Markdown exported to" in _tool_text(export_result)
    assert "<!-- wadocx:fidelity-bundle" not in editable_md_path.read_text(encoding="utf-8")
    assert "replaced from markdown" in _tool_text(import_result)
    assert len(Document(rebuilt_doc_path).inline_shapes) == 1
