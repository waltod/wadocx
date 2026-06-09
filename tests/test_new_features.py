import asyncio
import json
import zipfile
from pathlib import Path

import pytest
from docx import Document

from word_document_server.tools.content_tools import (
    add_caption,
    add_hyperlink,
    add_table_of_figures,
    get_document_statistics,
    set_page_setup,
)
from word_document_server.utils.document_utils import (
    find_and_replace_text,
    normalize_hex_color,
)


def _doc_xml(path: Path) -> str:
    with zipfile.ZipFile(path) as zf:
        return zf.read("word/document.xml").decode("utf-8", "ignore")


def _make_doc(path: Path) -> None:
    doc = Document()
    doc.add_heading("Intro", 1)
    doc.add_paragraph("Hello world, some words here.")
    doc.save(str(path))


# --- normalize_hex_color -----------------------------------------------------

def test_normalize_hex_color_variants():
    assert normalize_hex_color("#FF0000") == "FF0000"
    assert normalize_hex_color("ff0000") == "FF0000"
    assert normalize_hex_color("#f00") == "FF0000"  # shorthand expansion
    assert normalize_hex_color("red") == "FF0000"  # named
    assert normalize_hex_color(None) is None
    assert normalize_hex_color("") is None


def test_normalize_hex_color_rejects_garbage():
    with pytest.raises(ValueError):
        normalize_hex_color("not-a-color")
    with pytest.raises(ValueError):
        normalize_hex_color("12345")


# --- cross-run find and replace ---------------------------------------------

def test_find_and_replace_spans_runs(tmp_path: Path):
    path = tmp_path / "replace.docx"
    doc = Document()
    para = doc.add_paragraph()
    para.add_run("Hello ")
    para.add_run("wor")
    para.add_run("ld!")  # "Hello world!" split across three runs
    doc.save(str(path))

    reopened = Document(str(path))
    count = find_and_replace_text(reopened, "Hello world", "Goodbye earth")
    reopened.save(str(path))

    assert count == 1
    assert Document(str(path)).paragraphs[0].text == "Goodbye earth!"


# --- hyperlink ---------------------------------------------------------------

def test_add_external_hyperlink(tmp_path: Path):
    path = tmp_path / "link.docx"
    _make_doc(path)
    msg = asyncio.run(add_hyperlink(str(path), "https://example.com", "Example"))
    assert "Hyperlink" in msg
    with zipfile.ZipFile(path) as zf:
        rels = zf.read("word/_rels/document.xml.rels").decode("utf-8", "ignore")
    assert "example.com" in rels
    assert 'TargetMode="External"' in rels
    assert "<w:hyperlink" in _doc_xml(path)


# --- statistics --------------------------------------------------------------

def test_document_statistics(tmp_path: Path):
    path = tmp_path / "stats.docx"
    _make_doc(path)
    asyncio.run(add_hyperlink(str(path), "https://example.com", "Example"))
    stats = json.loads(asyncio.run(get_document_statistics(str(path))))
    assert stats["words"] >= 5
    assert stats["headings"] == 1
    assert stats["headings_by_level"]["Heading 1"] == 1
    assert stats["hyperlinks"] == 1
    assert stats["sections"] == 1


# --- page setup --------------------------------------------------------------

def test_set_page_setup_landscape_a4(tmp_path: Path):
    path = tmp_path / "page.docx"
    _make_doc(path)
    msg = asyncio.run(
        set_page_setup(str(path), orientation="landscape", page_size="a4", margin_left=1.0)
    )
    assert "Page setup updated" in msg
    section = Document(str(path)).sections[0]
    assert section.page_width > section.page_height  # landscape


def test_set_page_setup_validates(tmp_path: Path):
    path = tmp_path / "page2.docx"
    _make_doc(path)
    assert "Invalid orientation" in asyncio.run(
        set_page_setup(str(path), orientation="sideways")
    )
    assert "Invalid page_size" in asyncio.run(set_page_setup(str(path), page_size="a7"))


# --- captions and table of figures ------------------------------------------

def test_caption_and_table_of_figures(tmp_path: Path):
    path = tmp_path / "fig.docx"
    _make_doc(path)
    assert "caption added" in asyncio.run(
        add_caption(str(path), "First chart", label="Figure")
    )
    assert "Table of figures" in asyncio.run(
        add_table_of_figures(str(path), label="Figure")
    )
    xml = _doc_xml(path)
    assert "SEQ Figure" in xml
    assert 'TOC \\h \\z \\c "Figure"' in xml
    # Document still reloads cleanly.
    assert len(Document(str(path)).paragraphs) >= 3
