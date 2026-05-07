import asyncio
import zipfile
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK

from word_document_server.tools.content_tools import (
    add_bookmark_to_paragraph,
    add_internal_hyperlink,
    add_live_table_of_contents,
    insert_omml_equation,
    set_document_footer_page_number,
    set_document_header_page_number,
)
from word_document_server.tools.content_tools import add_heading, add_paragraph
from word_document_server.tools.document_tools import create_document
from word_document_server.utils.markdown_utils import document_to_markdown, replace_document_with_markdown


def _read_zip_part(doc_path: Path, part_name: str) -> str:
    with zipfile.ZipFile(doc_path) as zf:
        return zf.read(part_name).decode("utf-8", "ignore")


def test_add_live_table_of_contents_inserts_native_field(tmp_path: Path):
    doc_path = tmp_path / "toc-field.docx"
    asyncio.run(create_document(str(doc_path), title="TOC Field Test"))
    asyncio.run(add_heading(str(doc_path), "Scope", level=1))
    asyncio.run(add_heading(str(doc_path), "Details", level=2))

    message = asyncio.run(
        add_live_table_of_contents(
            str(doc_path),
            title="Contents",
            max_level=3,
            insert_at_start=True,
        )
    )

    assert "Live table of contents inserted" in message

    document_xml = _read_zip_part(doc_path, "word/document.xml")
    settings_xml = _read_zip_part(doc_path, "word/settings.xml")
    assert 'TOC \\o "1-3" \\h \\z \\u' in document_xml
    assert "<w:updateFields" in settings_xml


def test_markdown_toc_marker_inserts_native_field(tmp_path: Path):
    doc_path = tmp_path / "markdown-toc.docx"
    markdown = """# Cover

<!-- TOC -->

# Scope

## Details
"""

    result = replace_document_with_markdown(str(doc_path), markdown)

    assert result["inserted_blocks"] == 5

    document_xml = _read_zip_part(doc_path, "word/document.xml")
    settings_xml = _read_zip_part(doc_path, "word/settings.xml")
    assert 'TOC \\o "1-3" \\h \\z \\u' in document_xml
    assert "Right-click to update field." in document_xml
    assert "<w:updateFields" in settings_xml


def test_markdown_wadocx_toc_directive_supports_options(tmp_path: Path):
    doc_path = tmp_path / "markdown-configured-toc.docx"
    markdown = """<!-- wadocx:toc
title: DAFTAR ISI
max_level: 2
page_break_after: true
-->

# BAB I

### Skipped Level
"""

    result = replace_document_with_markdown(str(doc_path), markdown)

    assert result["inserted_blocks"] == 5

    document_xml = _read_zip_part(doc_path, "word/document.xml")
    settings_xml = _read_zip_part(doc_path, "word/settings.xml")
    assert "DAFTAR ISI" in document_xml
    assert 'TOC \\o "1-2" \\h \\z \\u' in document_xml
    assert 'w:type="page"' in document_xml
    assert "<w:updateFields" in settings_xml


def test_markdown_toc_directive_supports_google_docs_like_styles(tmp_path: Path):
    styles = {
        "dotted": 'TOC \\o "1-3" \\h \\z \\u',
        "page_numbers": 'TOC \\o "1-3" \\h \\z \\u \\p " "',
        "links": 'TOC \\o "1-3" \\h \\z \\u \\n "1-3"',
    }

    for style, expected_instruction in styles.items():
        doc_path = tmp_path / f"markdown-toc-{style}.docx"
        markdown = f"""<!-- wadocx:toc
title: Contents
max_level: 3
style: {style}
-->

# Scope
"""

        replace_document_with_markdown(str(doc_path), markdown)

        document_xml = _read_zip_part(doc_path, "word/document.xml")
        assert expected_instruction in document_xml


def test_markdown_export_emits_supported_directives(tmp_path: Path):
    doc_path = tmp_path / "export-directives.docx"
    doc = Document()
    centered = doc.add_paragraph("Centered paragraph")
    centered.alignment = WD_ALIGN_PARAGRAPH.CENTER
    breaker = doc.add_paragraph("")
    breaker.add_run().add_break(WD_BREAK.PAGE)
    doc.add_heading("Scope", level=1)
    doc.save(doc_path)

    asyncio.run(
        add_live_table_of_contents(
            str(doc_path),
            title="",
            max_level=2,
            toc_style="links",
            insert_at_start=False,
        )
    )

    markdown = document_to_markdown(str(doc_path), include_fidelity_bundle=False)

    assert '<div align="center">' in markdown
    assert "<!-- PAGE BREAK -->" in markdown
    assert "<!-- wadocx:toc" in markdown
    assert "max_level: 2" in markdown
    assert "style: links" in markdown


def test_markdown_export_preserves_toc_title_in_directive(tmp_path: Path):
    source_doc_path = tmp_path / "toc-title-source.docx"
    rebuilt_doc_path = tmp_path / "toc-title-rebuilt.docx"
    doc = Document()
    doc.add_heading("Scope", level=1)
    doc.save(source_doc_path)

    asyncio.run(
        add_live_table_of_contents(
            str(source_doc_path),
            title="Contents",
            max_level=2,
            insert_at_start=True,
        )
    )

    markdown = document_to_markdown(str(source_doc_path), include_fidelity_bundle=False)
    result = replace_document_with_markdown(str(rebuilt_doc_path), markdown)

    assert "<!-- wadocx:toc" in markdown
    assert "title: Contents" in markdown
    assert not markdown.startswith("Contents\n\n<!-- wadocx:toc")
    assert result["inserted_blocks"] >= 2
    rebuilt_xml = _read_zip_part(rebuilt_doc_path, "word/document.xml")
    assert "Contents" in rebuilt_xml
    assert 'TOC \\o "1-2" \\h \\z \\u' in rebuilt_xml


def test_markdown_toc_alignment_round_trips(tmp_path: Path):
    source_doc_path = tmp_path / "toc-alignment-source.docx"
    rebuilt_doc_path = tmp_path / "toc-alignment-rebuilt.docx"
    hand_authored_doc_path = tmp_path / "toc-alignment-hand-authored.docx"
    doc = Document()
    doc.add_heading("Scope", level=1)
    doc.save(source_doc_path)

    asyncio.run(
        add_live_table_of_contents(
            str(source_doc_path),
            title="Contents",
            max_level=2,
            insert_at_start=True,
        )
    )
    source = Document(source_doc_path)
    source.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    source.save(source_doc_path)

    markdown = document_to_markdown(str(source_doc_path), include_fidelity_bundle=False)
    replace_document_with_markdown(str(rebuilt_doc_path), markdown)
    hand_result = replace_document_with_markdown(
        str(hand_authored_doc_path),
        """<div align="center">
<!-- wadocx:toc
title: Contents
max_level: 2
-->
</div>

# Scope
""",
    )

    rebuilt = Document(rebuilt_doc_path)
    hand_authored = Document(hand_authored_doc_path)
    assert '<div align="center">' in markdown
    assert rebuilt.paragraphs[0].alignment == WD_ALIGN_PARAGRAPH.CENTER
    assert rebuilt.paragraphs[1].alignment == WD_ALIGN_PARAGRAPH.CENTER
    assert hand_result["inserted_blocks"] >= 2
    assert hand_authored.paragraphs[0].alignment == WD_ALIGN_PARAGRAPH.CENTER
    assert hand_authored.paragraphs[1].alignment == WD_ALIGN_PARAGRAPH.CENTER


def test_header_and_footer_page_fields_are_written_as_native_fields(tmp_path: Path):
    doc_path = tmp_path / "page-fields.docx"
    asyncio.run(create_document(str(doc_path), title="Page Field Test"))

    header_message = asyncio.run(
        set_document_header_page_number(
            str(doc_path),
            prefix_text="Page ",
            suffix_text=" draft",
            font_name="Arial",
            font_size=12,
        )
    )
    footer_message = asyncio.run(
        set_document_footer_page_number(
            str(doc_path),
            prefix_text="Page ",
            footer_type="even",
            alignment="center",
        )
    )

    assert "Header page number field updated" in header_message
    assert "Footer page number field updated" in footer_message

    header_xml = _read_zip_part(doc_path, "word/header1.xml")
    footer_xml = _read_zip_part(doc_path, "word/footer1.xml")
    settings_xml = _read_zip_part(doc_path, "word/settings.xml")

    assert " PAGE " in header_xml
    assert " PAGE " in footer_xml
    assert "<w:updateFields" in settings_xml


def test_insert_omml_equation_writes_math_object(tmp_path: Path):
    doc_path = tmp_path / "equation.docx"
    asyncio.run(create_document(str(doc_path), title="Equation Test"))
    asyncio.run(add_paragraph(str(doc_path), "Intro paragraph"))

    message = asyncio.run(
        insert_omml_equation(str(doc_path), "E = mc^2", paragraph_index=0, position="after")
    )
    assert "OMML equation inserted" in message

    document_xml = _read_zip_part(doc_path, "word/document.xml")
    assert "<m:oMathPara" in document_xml
    assert "E = mc^2" in document_xml


def test_bookmark_and_internal_hyperlink_are_written(tmp_path: Path):
    doc_path = tmp_path / "bookmark-hyperlink.docx"
    asyncio.run(create_document(str(doc_path), title="Bookmark Test"))
    asyncio.run(add_paragraph(str(doc_path), "Destination paragraph"))
    asyncio.run(add_paragraph(str(doc_path), "See destination: "))

    bookmark_message = asyncio.run(
        add_bookmark_to_paragraph(str(doc_path), 0, "dest_clause")
    )
    hyperlink_message = asyncio.run(
        add_internal_hyperlink(str(doc_path), 1, "jump", "dest_clause")
    )

    assert "Bookmark 'dest_clause' added" in bookmark_message
    assert "Internal hyperlink to 'dest_clause' added" in hyperlink_message

    document_xml = _read_zip_part(doc_path, "word/document.xml")
    assert 'w:bookmarkStart' in document_xml
    assert 'w:name="dest_clause"' in document_xml
    assert 'w:hyperlink' in document_xml
    assert 'w:anchor="dest_clause"' in document_xml
