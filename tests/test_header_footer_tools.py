import asyncio
from pathlib import Path

from docx import Document
from docx.shared import Pt

from word_document_server.tools.content_tools import (
    get_document_footer,
    get_document_header,
    set_document_footer,
    set_document_header,
)
from word_document_server.tools.document_tools import create_document


def test_set_and_get_default_header_footer(tmp_path: Path):
    doc_path = tmp_path / "header-footer.docx"
    asyncio.run(create_document(str(doc_path), title="Header Footer Test"))

    header_message = asyncio.run(
        set_document_header(
            str(doc_path),
            "Confidential Draft",
            font_name="Arial",
            font_size=14,
            bold=True,
            color="FF0000",
            alignment="center",
        )
    )
    footer_message = asyncio.run(
        set_document_footer(
            str(doc_path),
            "Page footer note",
            font_name="Calibri",
            font_size=10,
            italic=True,
            color="0000FF",
            alignment="right",
        )
    )

    assert "Header text updated" in header_message
    assert "Footer text updated" in footer_message
    assert asyncio.run(get_document_header(str(doc_path))) == "Confidential Draft"
    assert asyncio.run(get_document_footer(str(doc_path))) == "Page footer note"

    doc = Document(doc_path)
    header_para = doc.sections[0].header.paragraphs[0]
    footer_para = doc.sections[0].footer.paragraphs[0]

    assert header_para.text == "Confidential Draft"
    assert footer_para.text == "Page footer note"
    assert header_para.runs[0].font.name == "Arial"
    assert header_para.runs[0].font.size == Pt(14)
    assert header_para.runs[0].font.bold is True
    assert str(header_para.runs[0].font.color.rgb) == "FF0000"
    assert footer_para.runs[0].font.name == "Calibri"
    assert footer_para.runs[0].font.size == Pt(10)
    assert footer_para.runs[0].font.italic is True
    assert str(footer_para.runs[0].font.color.rgb) == "0000FF"


def test_set_first_page_header_and_even_page_footer(tmp_path: Path):
    doc_path = tmp_path / "special-header-footer.docx"
    asyncio.run(create_document(str(doc_path), title="Special Header Footer Test"))

    first_header = asyncio.run(
        set_document_header(str(doc_path), "First Page Header", header_type="first")
    )
    even_footer = asyncio.run(
        set_document_footer(str(doc_path), "Even Footer", footer_type="even")
    )

    assert "section 0 (first)" in first_header
    assert "section 0 (even)" in even_footer
    assert asyncio.run(get_document_header(str(doc_path), header_type="first")) == "First Page Header"
    assert asyncio.run(get_document_footer(str(doc_path), footer_type="even")) == "Even Footer"

    doc = Document(doc_path)
    assert doc.sections[0].different_first_page_header_footer is True
    assert doc.settings.odd_and_even_pages_header_footer is True
    assert doc.sections[0].first_page_header.paragraphs[0].text == "First Page Header"
    assert doc.sections[0].even_page_footer.paragraphs[0].text == "Even Footer"
