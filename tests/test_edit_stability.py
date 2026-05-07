import asyncio
import base64
import hashlib
import shutil
import zipfile
from pathlib import Path

from docx import Document
from docx.enum.text import WD_BREAK

from word_document_server.tools.format_tools import format_text
from word_document_server.tools.markdown_tools import (
    export_document_markdown,
    get_document_markdown,
    replace_document_with_markdown,
    replace_document_with_markdown_file,
    replace_section_with_markdown,
)
from word_document_server.utils.document_utils import (
    replace_block_between_manual_anchors,
    replace_paragraph_block_below_header,
)


def _paragraph_texts(doc_path: Path):
    return [paragraph.text for paragraph in Document(doc_path).paragraphs if paragraph.text]


def _write_png(path: Path) -> None:
    png_base64 = (
        "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mP8/x8AAwMCAO7Z0mQAAAAASUVORK5CYII="
    )
    path.write_bytes(base64.b64decode(png_base64))


def _document_xml(doc_path: Path) -> str:
    with zipfile.ZipFile(doc_path) as zf:
        return zf.read("word/document.xml").decode("utf-8", "ignore")


def test_replace_paragraph_block_below_header_is_stable_across_iterations(tmp_path: Path):
    doc_path = tmp_path / "iterative-header.docx"
    doc = Document()
    doc.add_heading("Scope", level=1)
    doc.add_paragraph("Old clause 1")
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Legacy"
    table.cell(0, 1).text = "Block"
    table.cell(1, 0).text = "Should"
    table.cell(1, 1).text = "Go"
    doc.add_paragraph("Old clause 2")
    doc.add_heading("Next Section", level=1)
    doc.add_paragraph("Keep me")
    doc.save(doc_path)

    for iteration in range(6):
        message = replace_paragraph_block_below_header(
            str(doc_path),
            "Scope",
            [f"Clause revision {iteration}", f"Subclause revision {iteration}"],
        )
        assert "Replaced content under 'Scope'" in message

    final_doc = Document(doc_path)
    assert [paragraph.text for paragraph in final_doc.paragraphs if paragraph.text] == [
        "Scope",
        "Clause revision 5",
        "Subclause revision 5",
        "Next Section",
        "Keep me",
    ]
    assert len(final_doc.tables) == 0


def test_replace_block_between_manual_anchors_is_stable_across_iterations(tmp_path: Path):
    doc_path = tmp_path / "iterative-anchors.docx"
    doc = Document()
    doc.add_paragraph("START")
    doc.add_paragraph("Legacy paragraph")
    table = doc.add_table(rows=1, cols=1)
    table.cell(0, 0).text = "Legacy table"
    doc.add_paragraph("Another legacy paragraph")
    doc.add_paragraph("END")
    doc.add_paragraph("Trailing content")
    doc.save(doc_path)

    for iteration in range(5):
        message = replace_block_between_manual_anchors(
            str(doc_path),
            "START",
            [f"Anchor revision {iteration}", f"Anchor notes {iteration}"],
            end_anchor_text="END",
        )
        assert "Replaced content between 'START' and 'END'" in message

    final_doc = Document(doc_path)
    assert [paragraph.text for paragraph in final_doc.paragraphs if paragraph.text] == [
        "START",
        "Anchor revision 4",
        "Anchor notes 4",
        "END",
        "Trailing content",
    ]
    assert len(final_doc.tables) == 0


def test_markdown_round_trip_and_section_replacement(tmp_path: Path):
    doc_path = tmp_path / "markdown-flow.docx"
    initial_markdown = """# Standard Draft

Intro paragraph.

## Scope

- Item A
- Item B

| Clause | Value |
| --- | --- |
| A1 | Draft |
"""

    replace_message = asyncio.run(replace_document_with_markdown(str(doc_path), initial_markdown))
    assert "replaced from markdown" in replace_message

    exported_markdown = asyncio.run(get_document_markdown(str(doc_path)))
    assert "# Standard Draft" in exported_markdown
    assert "## Scope" in exported_markdown
    assert "- Item A" in exported_markdown
    assert exported_markdown.count("- Item A") == 1
    assert exported_markdown.count("- Item B") == 1
    assert "| Clause | Value |" in exported_markdown

    section_markdown = """Scope paragraph updated.

### Clause Details

1. First clause
2. Second clause
"""

    section_message = asyncio.run(
        replace_section_with_markdown(str(doc_path), "Scope", section_markdown)
    )
    assert "Section 'Scope' replaced from markdown" in section_message

    final_markdown = asyncio.run(get_document_markdown(str(doc_path)))
    assert "Scope paragraph updated." in final_markdown
    assert "### Clause Details" in final_markdown
    assert "1. First clause" in final_markdown
    assert "2. Second clause" in final_markdown


def test_docx_native_lists_export_once(tmp_path: Path):
    doc_path = tmp_path / "native-lists.docx"
    doc = Document()
    doc.add_paragraph("Bullet A", style="List Bullet")
    doc.add_paragraph("Bullet B", style="List Bullet")
    doc.add_paragraph("Number A", style="List Number")
    doc.add_paragraph("Number B", style="List Number")
    doc.save(doc_path)

    exported_markdown = asyncio.run(
        get_document_markdown(str(doc_path), include_fidelity_bundle=False)
    )

    assert exported_markdown.count("- Bullet A") == 1
    assert exported_markdown.count("- Bullet B") == 1
    assert exported_markdown.count("1. Number A") == 1
    assert exported_markdown.count("2. Number B") == 1


def test_format_text_preserves_plain_text_content(tmp_path: Path):
    doc_path = tmp_path / "format-text.docx"
    doc = Document()
    doc.add_paragraph("Alpha Beta Gamma")
    doc.save(doc_path)

    message = asyncio.run(format_text(str(doc_path), 0, 6, 10, bold=True, color="red"))
    assert "formatted successfully" in message

    final_doc = Document(doc_path)
    assert final_doc.paragraphs[0].text == "Alpha Beta Gamma"


def test_markdown_image_round_trip(tmp_path: Path):
    doc_path = tmp_path / "markdown-image.docx"
    image_path = tmp_path / "risk-flow.png"
    _write_png(image_path)

    markdown = f"""# Visual Draft

![Figure 5 - Risk and degraded-mode decision flow for landslide screening support](<{image_path}>)

Image paragraph.
"""

    replace_message = asyncio.run(replace_document_with_markdown(str(doc_path), markdown))
    assert "replaced from markdown" in replace_message

    doc = Document(doc_path)
    assert len(doc.inline_shapes) == 1
    assert doc.paragraphs[0].text == "Visual Draft"
    assert doc.paragraphs[-1].text == "Image paragraph."

    exported_markdown = asyncio.run(get_document_markdown(str(doc_path)))
    assert "![Figure 5 - Risk and degraded-mode decision flow for landslide screening support]" in exported_markdown
    assert "Image paragraph." in exported_markdown


def test_markdown_fidelity_bundle_restores_exact_docx(tmp_path: Path):
    source_doc_path = tmp_path / "source.docx"
    restored_doc_path = tmp_path / "restored.docx"

    doc = Document()
    doc.add_heading("Scope", level=1)
    doc.add_paragraph("Exact round-trip text.")
    doc.add_table(rows=2, cols=2)
    doc.tables[0].cell(0, 0).text = "A"
    doc.tables[0].cell(0, 1).text = "B"
    doc.tables[0].cell(1, 0).text = "C"
    doc.tables[0].cell(1, 1).text = "D"
    doc.save(source_doc_path)

    exported_markdown = asyncio.run(get_document_markdown(str(source_doc_path)))
    assert "<!-- wadocx:fidelity-bundle" in exported_markdown
    assert "# Scope" in exported_markdown

    restore_message = asyncio.run(
        replace_document_with_markdown(str(restored_doc_path), exported_markdown)
    )
    assert "restored exactly from a wadocx fidelity bundle" in restore_message

    source_hash = hashlib.sha256(source_doc_path.read_bytes()).hexdigest()
    restored_hash = hashlib.sha256(restored_doc_path.read_bytes()).hexdigest()
    assert source_hash == restored_hash


def test_markdown_explicit_exact_restore_mode(tmp_path: Path):
    source_doc_path = tmp_path / "source.docx"
    restored_doc_path = tmp_path / "restored.docx"

    doc = Document()
    doc.add_heading("Explicit Restore", level=1)
    doc.add_paragraph("Exact mode content.")
    doc.save(source_doc_path)

    exported_markdown = asyncio.run(get_document_markdown(str(source_doc_path)))
    restore_message = asyncio.run(
        replace_document_with_markdown(
            str(restored_doc_path),
            exported_markdown,
            mode="exact_restore",
        )
    )
    missing_bundle_message = asyncio.run(
        replace_document_with_markdown(
            str(tmp_path / "missing-bundle.docx"),
            "# Editable only\n",
            mode="exact_restore",
        )
    )

    assert "restored exactly from a wadocx fidelity bundle" in restore_message
    assert source_doc_path.read_bytes() == restored_doc_path.read_bytes()
    assert "Exact restore mode requires a wadocx:fidelity-bundle" in missing_bundle_message


def test_markdown_file_exact_restore_tolerates_bom_and_leading_blank_lines(tmp_path: Path):
    source_doc_path = tmp_path / "source.docx"
    markdown_path = tmp_path / "source.md"
    restored_doc_path = tmp_path / "restored.docx"

    doc = Document()
    doc.add_heading("BOM Restore", level=1)
    doc.add_paragraph("Exact restore with leading whitespace.")
    doc.save(source_doc_path)

    exported_markdown = asyncio.run(get_document_markdown(str(source_doc_path)))
    markdown_path.write_text(f"\ufeff\n\n{exported_markdown}", encoding="utf-8")

    restore_message = asyncio.run(
        replace_document_with_markdown_file(
            str(restored_doc_path),
            str(markdown_path),
            mode="exact_restore",
        )
    )

    assert "restored exactly from a wadocx fidelity bundle" in restore_message
    assert source_doc_path.read_bytes() == restored_doc_path.read_bytes()


def test_markdown_export_tool_alias_writes_file_for_file_backed_round_trip(tmp_path: Path):
    source_doc_path = tmp_path / "source.docx"
    exported_md_path = tmp_path / "source-export.md"
    restored_doc_path = tmp_path / "restored.docx"

    doc = Document()
    doc.add_heading("Alias Export", level=1)
    doc.add_paragraph("File-backed exact restore.")
    doc.save(source_doc_path)

    export_message = asyncio.run(
        export_document_markdown(str(source_doc_path), str(exported_md_path))
    )
    restore_message = asyncio.run(
        replace_document_with_markdown_file(str(restored_doc_path), str(exported_md_path))
    )

    assert "Markdown exported to" in export_message
    assert exported_md_path.exists()
    assert "<!-- wadocx:fidelity-bundle" in exported_md_path.read_text(encoding="utf-8")
    assert "restored exactly from a wadocx fidelity bundle" in restore_message
    assert source_doc_path.read_bytes() == restored_doc_path.read_bytes()


def test_markdown_export_tool_alias_can_omit_fidelity_bundle_for_editable_file(tmp_path: Path):
    source_doc_path = tmp_path / "source.docx"
    exported_md_path = tmp_path / "editable-export.md"
    rebuilt_doc_path = tmp_path / "rebuilt.docx"

    doc = Document()
    doc.add_heading("Editable Alias Export", level=1)
    doc.add_paragraph("File-backed editable rebuild.")
    doc.save(source_doc_path)

    export_message = asyncio.run(
        export_document_markdown(
            str(source_doc_path),
            str(exported_md_path),
            include_fidelity_bundle=False,
        )
    )
    restore_message = asyncio.run(
        replace_document_with_markdown_file(str(rebuilt_doc_path), str(exported_md_path))
    )

    exported_markdown = exported_md_path.read_text(encoding="utf-8")
    assert "Markdown exported to" in export_message
    assert "<!-- wadocx:fidelity-bundle" not in exported_markdown
    assert "# Editable Alias Export" in exported_markdown
    assert "replaced from markdown" in restore_message
    assert "File-backed editable rebuild." in _paragraph_texts(rebuilt_doc_path)


def test_markdown_export_file_uses_portable_relative_image_links(tmp_path: Path):
    source_dir = tmp_path / "source-dir"
    export_dir = tmp_path / "exports"
    source_dir.mkdir()
    export_dir.mkdir()
    source_doc_path = source_dir / "source.docx"
    exported_md_path = export_dir / "custom-name.md"
    relocated_dir = tmp_path / "relocated"
    relocated_md_path = relocated_dir / "custom-name.md"
    rebuilt_doc_path = tmp_path / "rebuilt.docx"
    image_path = tmp_path / "risk-flow.png"
    _write_png(image_path)

    replace_message = asyncio.run(
        replace_document_with_markdown(
            str(source_doc_path),
            f"""# Image Export

![Risk flow](<{image_path}>)
""",
        )
    )
    export_message = asyncio.run(
        export_document_markdown(
            str(source_doc_path),
            str(exported_md_path),
            include_fidelity_bundle=False,
        )
    )

    exported_markdown = exported_md_path.read_text(encoding="utf-8")
    media_dir = export_dir / "custom-name_media"
    relocated_dir.mkdir()
    shutil.copy2(exported_md_path, relocated_md_path)
    shutil.copytree(media_dir, relocated_dir / media_dir.name)
    rebuild_message = asyncio.run(
        replace_document_with_markdown_file(str(rebuilt_doc_path), str(relocated_md_path))
    )

    assert "replaced from markdown" in replace_message
    assert "Markdown exported to" in export_message
    assert "![Risk flow](<custom-name_media/image1.png>)" in exported_markdown
    assert media_dir.exists()
    assert str(media_dir) not in exported_markdown
    assert "replaced from markdown" in rebuild_message
    assert len(Document(rebuilt_doc_path).inline_shapes) == 1


def test_markdown_export_can_omit_fidelity_bundle_for_editable_rebuild(tmp_path: Path):
    source_doc_path = tmp_path / "source.docx"
    rebuilt_doc_path = tmp_path / "rebuilt.docx"

    doc = Document()
    doc.add_heading("Scope", level=1)
    doc.add_paragraph("Editable text.")
    doc.save(source_doc_path)

    exported_markdown = asyncio.run(
        get_document_markdown(str(source_doc_path), include_fidelity_bundle=False)
    )

    assert "<!-- wadocx:fidelity-bundle" not in exported_markdown
    assert "# Scope" in exported_markdown

    message = asyncio.run(
        replace_document_with_markdown(str(rebuilt_doc_path), exported_markdown)
    )

    assert "replaced from markdown" in message
    assert "Editable text." in _paragraph_texts(rebuilt_doc_path)


def test_markdown_replace_mode_can_ignore_fidelity_bundle(tmp_path: Path):
    source_doc_path = tmp_path / "source.docx"
    rebuilt_doc_path = tmp_path / "rebuilt.docx"

    doc = Document()
    doc.add_heading("Original", level=1)
    doc.add_paragraph("Old text.")
    doc.save(source_doc_path)

    exported_markdown = asyncio.run(get_document_markdown(str(source_doc_path)))
    edited_markdown = exported_markdown + "\n\n# Added\n\nNew text."

    message = asyncio.run(
        replace_document_with_markdown(
            str(rebuilt_doc_path),
            edited_markdown,
            mode="editable_rebuild",
        )
    )

    assert "replaced from markdown" in message
    rebuilt_texts = _paragraph_texts(rebuilt_doc_path)
    assert "Added" in rebuilt_texts
    assert "New text." in rebuilt_texts


def test_markdown_replace_mode_can_use_fidelity_bundle_as_template(tmp_path: Path):
    source_doc_path = tmp_path / "source.docx"
    rebuilt_doc_path = tmp_path / "rebuilt.docx"

    doc = Document()
    doc.sections[0].header.paragraphs[0].text = "Template header"
    doc.add_paragraph("Template cover")
    doc.add_section()
    doc.sections[1].footer.paragraphs[0].text = "Template footer"
    doc.add_paragraph("Template body")
    doc.save(source_doc_path)

    exported_markdown = asyncio.run(get_document_markdown(str(source_doc_path)))
    bundle = exported_markdown.split("-->", 1)[0] + "-->"
    edited_markdown = f"""{bundle}

# Cover

New cover.

<!-- SECTION BREAK -->

# Body

New body.
"""

    message = asyncio.run(
        replace_document_with_markdown(
            str(rebuilt_doc_path),
            edited_markdown,
            mode="editable_rebuild_with_template",
        )
    )

    assert "replaced from markdown" in message
    rebuilt = Document(rebuilt_doc_path)
    assert len(rebuilt.sections) == 2
    assert rebuilt.sections[0].header.paragraphs[0].text == "Template header"
    assert rebuilt.sections[1].footer.paragraphs[0].text == "Template footer"
    assert "New body." in [p.text for p in rebuilt.paragraphs]


def test_markdown_export_emits_section_boundaries_for_template_rebuild(tmp_path: Path):
    source_doc_path = tmp_path / "source.docx"
    rebuilt_doc_path = tmp_path / "rebuilt.docx"

    doc = Document()
    doc.sections[0].header.paragraphs[0].text = "Section one header"
    doc.add_heading("Cover", level=1)
    doc.add_paragraph("Cover text")
    doc.add_section()
    doc.sections[1].footer.paragraphs[0].text = "Section two footer"
    doc.add_heading("Body", level=1)
    doc.add_paragraph("Body text")
    doc.save(source_doc_path)

    exported_markdown = asyncio.run(get_document_markdown(str(source_doc_path)))
    message = asyncio.run(
        replace_document_with_markdown(
            str(rebuilt_doc_path),
            exported_markdown,
            mode="editable_rebuild_with_template",
        )
    )

    assert "<!-- SECTION BREAK -->" in exported_markdown
    assert "replaced from markdown" in message
    rebuilt = Document(rebuilt_doc_path)
    assert len(rebuilt.sections) == 2
    assert rebuilt.sections[0].header.paragraphs[0].text == "Section one header"
    assert rebuilt.sections[1].footer.paragraphs[0].text == "Section two footer"
    assert "Body text" in [p.text for p in rebuilt.paragraphs]


def test_markdown_template_rebuild_distinguishes_page_and_section_breaks(tmp_path: Path):
    source_doc_path = tmp_path / "mixed-breaks-source.docx"
    rebuilt_doc_path = tmp_path / "mixed-breaks-rebuilt.docx"

    doc = Document()
    doc.sections[0].header.paragraphs[0].text = "Section one header"
    doc.add_paragraph("Cover text")
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
    doc.add_paragraph("Middle text")
    doc.add_section()
    doc.sections[1].footer.paragraphs[0].text = "Section two footer"
    doc.add_paragraph("Body text")
    doc.save(source_doc_path)

    exported_markdown = asyncio.run(get_document_markdown(str(source_doc_path)))
    message = asyncio.run(
        replace_document_with_markdown(
            str(rebuilt_doc_path),
            exported_markdown,
            mode="editable_rebuild_with_template",
        )
    )

    assert exported_markdown.index("<!-- PAGE BREAK -->") < exported_markdown.index(
        "<!-- SECTION BREAK -->"
    )
    assert "replaced from markdown" in message
    rebuilt = Document(rebuilt_doc_path)
    rebuilt_xml = _document_xml(rebuilt_doc_path)
    assert len(rebuilt.sections) == 2
    assert rebuilt.sections[0].header.paragraphs[0].text == "Section one header"
    assert rebuilt.sections[1].footer.paragraphs[0].text == "Section two footer"
    assert rebuilt_xml.index("Middle text") < rebuilt_xml.index("<w:sectPr") < rebuilt_xml.index("Body text")


def test_section_replacement_rejects_fidelity_bundle(tmp_path: Path):
    doc_path = tmp_path / "section-bundle.docx"
    doc = Document()
    doc.add_heading("Scope", level=1)
    doc.add_paragraph("Original content")
    doc.save(doc_path)

    exported_markdown = asyncio.run(get_document_markdown(str(doc_path)))
    message = asyncio.run(
        replace_section_with_markdown(str(doc_path), "Scope", exported_markdown)
    )
    assert "does not accept a wadocx fidelity bundle" in message


def test_markdown_can_use_base_template_md_with_page_breaks(tmp_path: Path):
    template_doc_path = tmp_path / "template.docx"
    template_md_path = tmp_path / "template.md"
    output_doc_path = tmp_path / "derived.docx"

    template = Document()
    template.sections[0].header.paragraphs[0].text = "Template header"
    template.add_paragraph("Template cover")
    template.add_section()
    template.sections[1].footer.paragraphs[0].text = "Template footer"
    template.add_paragraph("Template body")
    template.save(template_doc_path)

    exported_template_markdown = asyncio.run(get_document_markdown(str(template_doc_path)))
    template_md_path.write_text(exported_template_markdown, encoding="utf-8")

    derived_markdown = f"""<!-- wadocx:base-template-md
path: {template_md_path}
-->

# Cover

Cover text.

<!-- SECTION BREAK -->

# Body

Body text.
"""

    replace_message = asyncio.run(
        replace_document_with_markdown(str(output_doc_path), derived_markdown)
    )
    assert "replaced from markdown" in replace_message

    derived = Document(output_doc_path)
    assert len(derived.sections) == 2
    assert derived.sections[0].header.paragraphs[0].text == "Template header"
    assert derived.sections[1].footer.paragraphs[0].text == "Template footer"
    assert "Cover" in [p.text for p in derived.paragraphs]
    assert "Body" in [p.text for p in derived.paragraphs]


def test_markdown_file_replacement_resolves_relative_image_paths(tmp_path: Path):
    output_doc_path = tmp_path / "relative-image.docx"
    markdown_path = tmp_path / "draft.md"
    image_path = tmp_path / "risk-flow.png"
    _write_png(image_path)

    markdown_path.write_text(
        """# Visual Draft

![Risk flow](<risk-flow.png>)
""",
        encoding="utf-8",
    )

    message = asyncio.run(
        replace_document_with_markdown_file(str(output_doc_path), str(markdown_path))
    )

    assert "replaced from markdown" in message
    assert len(Document(output_doc_path).inline_shapes) == 1


def test_markdown_file_replacement_resolves_relative_base_template(tmp_path: Path):
    template_doc_path = tmp_path / "template.docx"
    template_md_path = tmp_path / "template.md"
    draft_md_path = tmp_path / "draft.md"
    output_doc_path = tmp_path / "derived.docx"

    template = Document()
    template.sections[0].header.paragraphs[0].text = "Template header"
    template.add_paragraph("Template cover")
    template.save(template_doc_path)

    exported_template_markdown = asyncio.run(get_document_markdown(str(template_doc_path)))
    template_md_path.write_text(exported_template_markdown, encoding="utf-8")
    draft_md_path.write_text(
        """<!-- wadocx:base-template-md
path: template.md
-->

# Replacement

Body from relative template.
""",
        encoding="utf-8",
    )

    message = asyncio.run(
        replace_document_with_markdown_file(str(output_doc_path), str(draft_md_path))
    )

    assert "replaced from markdown" in message
    derived = Document(output_doc_path)
    assert derived.sections[0].header.paragraphs[0].text == "Template header"
    assert "Body from relative template." in [p.text for p in derived.paragraphs]


def test_markdown_file_base_template_tolerates_bom_and_leading_blank_lines(tmp_path: Path):
    template_doc_path = tmp_path / "template.docx"
    template_md_path = tmp_path / "template.md"
    draft_md_path = tmp_path / "draft.md"
    output_doc_path = tmp_path / "derived.docx"

    template = Document()
    template.sections[0].header.paragraphs[0].text = "Template header"
    template.add_paragraph("Template cover")
    template.save(template_doc_path)

    exported_template_markdown = asyncio.run(get_document_markdown(str(template_doc_path)))
    template_md_path.write_text(exported_template_markdown, encoding="utf-8")
    draft_md_path.write_text(
        """\ufeff

<!-- wadocx:base-template-md
path: template.md
-->

# Replacement

Body from whitespace-prefixed template directive.
""",
        encoding="utf-8",
    )

    message = asyncio.run(
        replace_document_with_markdown_file(str(output_doc_path), str(draft_md_path))
    )

    assert "replaced from markdown" in message
    derived = Document(output_doc_path)
    assert derived.sections[0].header.paragraphs[0].text == "Template header"
    assert "Body from whitespace-prefixed template directive." in [
        p.text for p in derived.paragraphs
    ]


def test_markdown_base_template_md_requires_fidelity_bundle(tmp_path: Path):
    output_doc_path = tmp_path / "derived.docx"
    template_md_path = tmp_path / "template-no-bundle.md"
    draft_md_path = tmp_path / "draft.md"
    template_md_path.write_text("# Template without bundle\n", encoding="utf-8")
    draft_markdown = """<!-- wadocx:base-template-md
path: template-no-bundle.md
-->

# Replacement
"""
    draft_md_path.write_text(draft_markdown, encoding="utf-8")

    text_message = asyncio.run(
        replace_document_with_markdown(
            str(output_doc_path),
            draft_markdown,
            source_base_dir=str(tmp_path),
        )
    )
    file_message = asyncio.run(
        replace_document_with_markdown_file(str(output_doc_path), str(draft_md_path))
    )

    assert "Base template markdown must include a wadocx:fidelity-bundle" in text_message
    assert "Base template markdown must include a wadocx:fidelity-bundle" in file_message


def test_markdown_import_errors_are_returned_as_tool_messages(tmp_path: Path):
    doc_path = tmp_path / "error.docx"
    markdown_path = tmp_path / "invalid-mode.md"
    markdown_path.write_text("# Body\n", encoding="utf-8")
    missing_template_markdown = """<!-- wadocx:base-template-md
path: missing-template.md
-->

# Body
"""
    missing_image_markdown = """# Visual

![Missing](<missing.png>)
"""

    template_message = asyncio.run(
        replace_document_with_markdown(
            str(doc_path),
            missing_template_markdown,
            source_base_dir=str(tmp_path),
        )
    )
    image_message = asyncio.run(
        replace_document_with_markdown(
            str(doc_path),
            missing_image_markdown,
            source_base_dir=str(tmp_path),
        )
    )
    invalid_mode_message = asyncio.run(
        replace_document_with_markdown(
            str(doc_path),
            "# Body\n",
            mode="not-a-real-mode",
        )
    )
    invalid_file_mode_message = asyncio.run(
        replace_document_with_markdown_file(
            str(doc_path),
            str(markdown_path),
            mode="not-a-real-mode",
        )
    )

    assert "Base template markdown does not exist" in template_message
    assert "Markdown image file not found" in image_message
    assert "Unsupported markdown replace mode" in invalid_mode_message
    assert "Supported modes:" in invalid_mode_message
    assert "Unsupported markdown replace mode" in invalid_file_mode_message
