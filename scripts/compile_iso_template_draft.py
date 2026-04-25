"""
Compile a draft markdown file into the ISO Word template format.
"""
from __future__ import annotations

import os
import re
import shutil
import sys
from typing import Any, Dict, List, Optional, Tuple

REPO_ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), ".."))
if REPO_ROOT not in sys.path:
    sys.path.insert(0, REPO_ROOT)

from docx import Document
from docx.enum.text import WD_BREAK
from docx.oxml.ns import qn

from word_document_server.utils.document_utils import (
    get_body_elements,
    get_paragraph_from_element,
    insert_content_blocks_after_element,
    is_paragraph_element,
    normalize_paragraph_text,
    remove_body_elements,
)
from word_document_server.utils.markdown_utils import parse_markdown_blocks


PAGE_BREAK_SPLIT_RE = re.compile(
    r"^\s*<!--\s*(?:PAGE BREAK|ODD PAGE BREAK; PAGE NUMBER RESTARTS AT 1)\s*-->\s*$",
    re.MULTILINE,
)


def split_segments(markdown_text: str) -> List[str]:
    text = re.sub(
        r"^\s*<!--\s*wadocx:base-template-md\s*\n.*?\n-->\s*",
        "",
        markdown_text,
        flags=re.DOTALL,
    ).strip()
    return [segment.strip() for segment in PAGE_BREAK_SPLIT_RE.split(text)]


def clear_paragraph_runs(paragraph) -> None:
    for run in list(paragraph.runs):
        paragraph._p.remove(run._r)


def add_run_like(paragraph, text: str, template_run) -> None:
    run = paragraph.add_run(text)
    if template_run is not None:
        run.bold = template_run.bold
        run.italic = template_run.italic
        run.underline = template_run.underline
        run.font.name = template_run.font.name
        run.font.size = template_run.font.size
        if template_run.font.color is not None and template_run.font.color.rgb is not None:
            run.font.color.rgb = template_run.font.color.rgb


def set_single_run_text(paragraph, text: str) -> None:
    template_run = paragraph.runs[0] if paragraph.runs else None
    clear_paragraph_runs(paragraph)
    add_run_like(paragraph, text, template_run)


def set_title_like(paragraph, text: str) -> None:
    template_runs = list(paragraph.runs)
    clear_paragraph_runs(paragraph)
    first_word, _, remainder = text.partition(" ")
    if template_runs:
        add_run_like(paragraph, first_word, template_runs[0])
        if remainder:
            add_run_like(paragraph, " " + remainder, template_runs[1] if len(template_runs) > 1 else template_runs[0])
    else:
        paragraph.add_run(text)


def parse_cover_segment(segment: str) -> Dict[str, Any]:
    right_block = re.search(r'<div\s+align="right">\s*(.*?)\s*</div>', segment, re.DOTALL | re.IGNORECASE)
    center_blocks = re.findall(r'<div\s+align="center">\s*(.*?)\s*</div>', segment, re.DOTALL | re.IGNORECASE)
    cleaned = re.sub(r'<div\s+align="(?:right|center)">\s*.*?\s*</div>', "", segment, flags=re.DOTALL | re.IGNORECASE)
    paragraphs = [block["text"] for block in parse_markdown_blocks(cleaned) if block.get("type") == "paragraph"]

    right_lines: List[str] = []
    if right_block:
        for raw_line in right_block.group(1).splitlines():
            line = raw_line.strip()
            if not line:
                continue
            line = re.sub(r"\*\*(.+?)\*\*", r"\1", line)
            line = line.rstrip("  ")
            right_lines.append(line)

    center_texts: List[str] = []
    for block in center_blocks:
        center_texts.extend(
            item["text"]
            for item in parse_markdown_blocks(block)
            if item.get("type") == "paragraph"
        )

    return {
        "right_lines": right_lines,
        "title": paragraphs[0] if paragraphs else "",
        "warning_paragraphs": paragraphs[1:],
        "stage": center_texts[0] if center_texts else "",
        "warning_title": center_texts[1] if len(center_texts) > 1 else "Warning for WDs and CDs",
    }


def make_paragraph_blocks(text: str, style: str = "Body Text") -> List[Dict[str, Any]]:
    return [
        {"type": "paragraph", "text": block["text"], "style": style}
        for block in parse_markdown_blocks(text)
        if block.get("type") == "paragraph"
    ]


def strip_clause_prefix(text: str) -> str:
    return re.sub(r"^[A-Z]?\d+(?:\.\d+)*\s+", "", text).strip()


def format_annex_heading(text: str) -> str:
    match = re.match(r"^Annex\s+([A-Z])\s+\(([^)]+)\)\s+(.+)$", text)
    if not match:
        return text
    _, annex_type, title = match.groups()
    return f"\n({annex_type})\n\n{title}"


def transform_body_blocks(blocks: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
    styled: List[Dict[str, Any]] = []
    current_clause = ""
    in_annex = False
    awaiting_term_name = False

    i = 0
    while i < len(blocks):
        block = blocks[i]
        block_type = block.get("type")

        if block_type == "heading":
            text = block.get("text", "").strip()
            level = int(block.get("level", 1))
            awaiting_term_name = False

            if level == 1 and text.lower().startswith("annex "):
                in_annex = True
                current_clause = text
                styled.append({"type": "paragraph", "text": format_annex_heading(text), "style": "ANNEX"})
            elif level == 1 and text.lower() == "bibliography":
                in_annex = False
                current_clause = text
                styled.append({"type": "paragraph", "text": "Bibliography", "style": "Biblio Title"})
            elif level == 2:
                current_clause = text
                style = "a2" if in_annex else "Heading 1"
                styled.append({"type": "paragraph", "text": strip_clause_prefix(text), "style": style})
            elif level == 3:
                style = "a3" if in_annex else "Heading 2"
                styled.append({"type": "paragraph", "text": strip_clause_prefix(text), "style": style})
            elif level == 4:
                style = "a4" if in_annex else "Heading 3"
                styled.append({"type": "paragraph", "text": strip_clause_prefix(text), "style": style})
            else:
                styled.append({"type": "paragraph", "text": text, "style": "Body Text"})
            i += 1
            continue

        if block_type == "paragraph":
            text = block.get("text", "").strip()
            if not text:
                i += 1
                continue

            if current_clause.lower().startswith("3 terms and definitions") and re.fullmatch(r"\d+\.\d+", text):
                styled.append({"type": "paragraph", "text": text, "style": "TermNum"})
                awaiting_term_name = True
                i += 1
                continue

            if awaiting_term_name:
                styled.append({"type": "paragraph", "text": text, "style": "Term(s)"})
                awaiting_term_name = False
                i += 1
                continue

            if text.startswith("Figure "):
                styled.append({"type": "paragraph", "text": text, "style": "Figure Title"})
                i += 1
                continue

            if text.startswith("Table "):
                styled.append({"type": "paragraph", "text": text, "style": "Table title"})
                i += 1
                continue

            style = "Definition" if current_clause.lower().startswith("3 terms and definitions") else "Body Text"
            styled.append({"type": "paragraph", "text": text, "style": style})
            i += 1
            continue

        if block_type == "image":
            image_block = dict(block)
            image_block["alignment"] = "center"
            image_block["width"] = 6.0
            styled.append(image_block)
            i += 1
            continue

        if block_type == "table":
            styled.append(block)
            i += 1
            continue

        if block_type == "list":
            items = block.get("items", [])
            if block.get("ordered"):
                for index, item in enumerate(items, start=1):
                    styled.append({"type": "paragraph", "text": f"{index}. {item}", "style": "Body Text"})
            else:
                for item in items:
                    styled.append({"type": "paragraph", "text": f"• {item}", "style": "Body Text"})
            i += 1
            continue

        i += 1

    return styled


def find_first_paragraph_by_style(doc, style_name: str):
    for para in doc.paragraphs:
        if para.style and para.style.name == style_name:
            return para
    raise ValueError(f"Paragraph with style '{style_name}' not found.")


def remove_after_paragraph(doc, paragraph) -> None:
    elements = get_body_elements(doc)
    target_idx = None
    target_element = paragraph._element
    for i, el in enumerate(elements):
        if is_paragraph_element(el):
            if el == target_element:
                target_idx = i
                break
    if target_idx is None:
        raise ValueError("Target paragraph not found in body.")
    remove_body_elements(doc, target_idx + 1, len(elements))


def replace_body_between_text_markers(
    doc,
    start_text: str,
    end_text: Optional[str],
    blocks: List[Dict[str, Any]],
    end_style: Optional[str] = None,
) -> None:
    elements = get_body_elements(doc)
    start_idx = None
    end_idx = len(elements)
    start_normalized = normalize_paragraph_text(start_text).lower()
    end_normalized = normalize_paragraph_text(end_text).lower() if end_text else None

    for i, el in enumerate(elements):
        if not is_paragraph_element(el):
            continue
        para = get_paragraph_from_element(doc, el)
        if para is None:
            continue
        if normalize_paragraph_text(para.text).lower() == start_normalized:
            start_idx = i
            break

    if start_idx is None:
        raise ValueError(f"Heading '{start_text}' not found.")

    if end_style:
        for i in range(start_idx + 1, len(elements)):
            if not is_paragraph_element(elements[i]):
                continue
            para = get_paragraph_from_element(doc, elements[i])
            if para is None or para.style is None:
                continue
            if para.style.name == end_style:
                end_idx = i
                break
    elif end_normalized:
        for i in range(start_idx + 1, len(elements)):
            if not is_paragraph_element(elements[i]):
                continue
            para = get_paragraph_from_element(doc, elements[i])
            if para is None:
                continue
            if normalize_paragraph_text(para.text).lower() == end_normalized:
                end_idx = i
                break

    protected_end_idx = end_idx
    if end_idx > start_idx + 1:
        previous_el = elements[end_idx - 1]
        if is_paragraph_element(previous_el):
            p_pr = previous_el.find(qn("w:pPr"))
            sect_pr = p_pr.find(qn("w:sectPr")) if p_pr is not None else None
            if sect_pr is not None:
                protected_end_idx = end_idx - 1

    remove_body_elements(doc, start_idx + 1, protected_end_idx)
    refreshed_elements = get_body_elements(doc)
    header_el = refreshed_elements[start_idx]
    insert_content_blocks_after_element(doc, header_el, blocks, default_paragraph_style="Body Text")


def ensure_page_break_before_text(doc, heading_text: str) -> None:
    normalized = normalize_paragraph_text(heading_text).lower()
    for para in doc.paragraphs:
        if normalize_paragraph_text(para.text).lower() == normalized:
            breaker = doc.add_paragraph("")
            breaker.add_run().add_break(WD_BREAK.PAGE)
            para._element.addprevious(breaker._element)
            return
    raise ValueError(f"Heading '{heading_text}' not found for page-break insertion.")


def compile_iso_draft(markdown_path: str, template_docx_path: str, output_docx_path: str) -> str:
    with open(markdown_path, "r", encoding="utf-8") as markdown_file:
        markdown_text = markdown_file.read()

    segments = split_segments(markdown_text)
    if len(segments) < 6:
        raise ValueError(f"Expected at least 6 markdown segments, found {len(segments)}.")

    cover = parse_cover_segment(segments[0])
    foreword_text = segments[3]
    introduction_text = segments[4]
    body_text = segments[5]

    shutil.copyfile(template_docx_path, output_docx_path)
    doc = Document(output_docx_path)

    right_lines = cover["right_lines"] + ["", "", ""]
    set_single_run_text(doc.paragraphs[0], right_lines[0])
    set_single_run_text(doc.paragraphs[1], right_lines[1])
    set_single_run_text(doc.paragraphs[2], right_lines[2])
    set_title_like(doc.paragraphs[3], cover["title"])
    set_single_run_text(doc.paragraphs[5], cover["stage"] or "WD stage")
    set_single_run_text(doc.paragraphs[7], cover["warning_title"])
    warning_paragraphs = cover["warning_paragraphs"]
    if warning_paragraphs:
        set_single_run_text(doc.paragraphs[8], warning_paragraphs[0])
    if len(warning_paragraphs) > 1:
        set_single_run_text(doc.paragraphs[9], warning_paragraphs[1])

    replace_body_between_text_markers(doc, "Foreword", "Introduction", make_paragraph_blocks(foreword_text))
    replace_body_between_text_markers(
        doc,
        "Introduction",
        None,
        make_paragraph_blocks(introduction_text),
        end_style="zzSTDTitle",
    )
    ensure_page_break_before_text(doc, "Introduction")

    body_blocks = parse_markdown_blocks(body_text)
    title_heading = next(
        (block["text"] for block in body_blocks if block.get("type") == "heading" and int(block.get("level", 1)) == 1),
        "",
    )
    if not title_heading:
        raise ValueError("Main title heading not found in body segment.")
    title_para = find_first_paragraph_by_style(doc, "zzSTDTitle")
    set_title_like(title_para, title_heading)
    remaining_body_blocks = []
    title_consumed = False
    for block in body_blocks:
        if not title_consumed and block.get("type") == "heading" and int(block.get("level", 1)) == 1:
            title_consumed = True
            continue
        remaining_body_blocks.append(block)

    remove_after_paragraph(doc, title_para)
    styled_blocks = transform_body_blocks(remaining_body_blocks)
    insert_content_blocks_after_element(doc, title_para._element, styled_blocks, default_paragraph_style="Body Text")
    doc.save(output_docx_path)
    return output_docx_path


def main() -> None:
    desktop_root = os.path.abspath(os.path.join(REPO_ROOT, "..", ".."))
    markdown_path = os.path.join(desktop_root, "KSN", "ksn2026", "Idea3_Landslide_Screening", "new_draft_formatted.md")
    template_docx_path = os.path.join(desktop_root, "KSN", "ksn2026", "format", "iso format.docx")
    output_docx_path = os.path.join(desktop_root, "KSN", "ksn2026", "Idea3_Landslide_Screening", "formatted_iso_actual.docx")
    print(compile_iso_draft(markdown_path, template_docx_path, output_docx_path))


if __name__ == "__main__":
    main()
